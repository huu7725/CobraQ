from fastapi import APIRouter, HTTPException, Header, Depends, UploadFile, File
from pydantic import BaseModel
import json, random, re, base64, io, traceback, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
sys.stderr.reconfigure(encoding="utf-8", errors="replace")
from datetime import datetime
from pathlib import Path
from typing import Optional

from ..core.security import get_current_user_optional, Role
from ..core.audit import audit_log, EventType
from ..services.evaluation import EvaluationEntry, evaluation_logger

router = APIRouter(prefix="/files", tags=["files"])


def resolve_uid(current_user, x_user_id: str):
    """Teacher/Admin có thể truy cập file của mọi user."""
    if current_user and current_user.get("role") in (Role.TEACHER.value, Role.ADMIN.value):
        return current_user.get("sub")  # Teacher/admin dùng user_id của mình
    return (current_user.get("sub") if current_user else None) or x_user_id


def _resolve_file_uid(current_user, x_user_id: str, x_file_for: str = ""):
    """Resolve which user's files to access. Teacher/Admin can access any user's files."""
    is_privileged = current_user and current_user.get("role") in (Role.TEACHER.value, Role.ADMIN.value)
    if is_privileged and x_file_for:
        return x_file_for
    if is_privileged:
        return current_user.get("sub")
    return (current_user.get("sub") if current_user else None) or x_user_id


def user_dir(uid):
    d = Path("data/users") / re.sub(r'[^\w]', '_', uid or "guest")
    d.mkdir(parents=True, exist_ok=True)
    return d


def load_json(path, default):
    try:
        if Path(path).exists():
            with open(path, encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return default


def save_json(path, data):
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def files_index_path(uid): return user_dir(uid) / "files_index.json"


# ── PARSERS (giữ nguyên từ main_updated.py) ──

def is_highlight(fill):
    if not fill or len(fill) < 3:
        return False
    r, g, b = float(fill[0]), float(fill[1]), float(fill[2])
    if r > 0.92 and g > 0.92 and b > 0.92: return False
    if r < 0.05 and g < 0.05 and b < 0.05: return False
    if r > 0.85 and g > 0.85 and b > 0.85: return False
    if g > 0.50 and b > 0.40 and r < 0.60: return True
    if g > 0.55 and g > r * 1.3 and g > b * 1.2: return True
    if r > 0.75 and g > 0.70 and b < 0.40: return True
    if r > 0.80 and g < 0.65 and b < 0.65: return True
    if b > 0.55 and b > r * 1.2 and g > 0.45: return True
    return False


def get_highlight_rects(page):
    rects = []
    try:
        for d in page.get_drawings():
            if is_highlight(d.get("fill")):
                rects.append(d["rect"])
    except:
        pass
    return rects


def span_in_rect(bbox, rects, pad=4):
    cx = (bbox[0] + bbox[2]) / 2
    cy = (bbox[1] + bbox[3]) / 2
    for r in rects:
        if r.x0-pad <= cx <= r.x1+pad and r.y0-pad <= cy <= r.y1+pad:
            return True
    return False


def get_all_spans(page):
    spans = []
    try:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                for sp in line["spans"]:
                    t = sp["text"].strip()
                    if t:
                        bbox = sp["bbox"]
                        spans.append({"text": t, "bbox": bbox,
                                     "cx": (bbox[0]+bbox[2])/2, "cy": (bbox[1]+bbox[3])/2})
    except:
        pass
    return spans


def group_rows(spans, tol=12):
    rows = []
    for sp in sorted(spans, key=lambda s: s["cy"]):
        placed = False
        for row in rows:
            if abs(row["cy"] - sp["cy"]) < tol:
                row["spans"].append(sp)
                row["cy"] = (row["cy"] + sp["cy"]) / 2
                placed = True
                break
        if not placed:
            rows.append({"cy": sp["cy"], "spans": [sp]})
    return sorted(rows, key=lambda r: r["cy"])


def parse_pdf_table(page, hl_rects):
    spans = get_all_spans(page)
    if not spans:
        return None
    rows = group_rows(spans)
    col_A = col_B = col_C = col_D = col_Q = None
    header_y = None

    for row in rows:
        has_da = bool(re.search(r'ĐÁPÁN|ANSWER', " ".join(s["text"] for s in row["spans"]), re.IGNORECASE))
        labels = [s for s in row["spans"] if s["text"].strip().upper() in ["A", "B", "C", "D"]]
        if has_da or len(labels) >= 3:
            header_y = row["cy"]
            for sp in row["spans"]:
                t = sp["text"].upper().replace(" ", "")
                if any(k in t for k in ["NỘIDUNG", "CÂUHỎI", "CONTENT"]):
                    col_Q = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]NA|ANSWERA', t) or t == "A": col_A = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]NB|ANSWERB', t) or t == "B": col_B = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]NC|ANSWERC', t) or t == "C": col_C = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]ND|ANSWERD', t) or t == "D": col_D = sp["cx"]
            break

    found_cols = sum(1 for c in [col_A, col_B, col_C, col_D] if c is not None)
    if found_cols < 2:
        return None

    col_map = {}
    for label, cx in [("A", col_A), ("B", col_B), ("C", col_C), ("D", col_D)]:
        if cx is not None:
            col_map[label] = cx

    answer_map = {}
    for hl in hl_rects:
        hcx = (hl.x0 + hl.x1) / 2
        hcy = (hl.y0 + hl.y1) / 2
        if header_y and hcy <= header_y + 5:
            continue
        best_label, best_dist = None, 999
        for label, cx in col_map.items():
            d = abs(hcx - cx)
            if d < best_dist:
                best_dist = d
                best_label = label
        if not best_label or best_dist > 120:
            continue
        best_tt, best_dy = None, 999
        for sp in spans:
            if re.match(r'^\d+$', sp["text"]):
                dy = abs(sp["cy"] - hcy)
                if dy < best_dy:
                    best_dy = dy
                    best_tt = int(sp["text"])
        if best_tt and best_dy < 30:
            answer_map[best_tt] = best_label

    if not answer_map:
        return None

    questions = []
    data_rows = [r for r in rows if header_y is None or r["cy"] > header_y + 8]
    tt_to_spans = {}
    current_tt = None
    for row in data_rows:
        tt_spans = [s for s in row["spans"] if re.match(r'^\d+$', s["text"])]
        if tt_spans:
            current_tt = int(tt_spans[0]["text"])
        if current_tt:
            if current_tt not in tt_to_spans:
                tt_to_spans[current_tt] = []
            tt_to_spans[current_tt].extend(row["spans"])

    for tt, row_spans in tt_to_spans.items():
        if tt not in answer_map:
            continue
        if col_Q:
            q_spans = [s for s in row_spans
                       if col_Q - 180 <= s["cx"] <= col_Q + 180
                       and s["text"] not in [str(tt)]]
        else:
            q_spans = [s for s in row_spans
                       if len(s["text"]) > 2
                       and not re.match(r'^[\dABCD]$', s["text"])
                       and not any(abs(s["cx"] - cx) < 80 for cx in col_map.values())]
        q_text = " ".join(s["text"] for s in sorted(q_spans, key=lambda x: (x["cy"], x["cx"])))
        if not q_text.strip():
            continue
        choices = []
        for label in sorted(col_map.keys()):
            cx = col_map[label]
            c_spans = [s for s in row_spans
                       if abs(s["cx"] - cx) < 90 and s["text"] not in [str(tt)]]
            c_text = " ".join(s["text"] for s in sorted(c_spans, key=lambda x: (x["cy"], x["cx"])))
            if c_text.strip():
                choices.append({"label": label, "text": c_text.strip()})
        if len(choices) >= 2:
            questions.append({
                "id": tt, "question": q_text.strip(),
                "choices": choices, "answer": answer_map.get(tt, ""),
                "explanation": "",
            })
    return questions if questions else None


def is_red_text(color_int):
    if color_int is None or color_int == 0:
        return False
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8)  & 0xFF
    b =  color_int        & 0xFF
    if r > 150 and g < 100 and b < 100: return True
    if r > 180 and g < 80 and b < 80: return True
    return False


def parse_pdf_inline(page, hl_rects):
    lines = []
    try:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                line_text = ""
                hl_bg = False
                hl_red = False
                for sp in line["spans"]:
                    line_text += sp["text"]
                    if hl_rects and span_in_rect(sp["bbox"], hl_rects):
                        hl_bg = True
                    if is_red_text(sp.get("color", 0)):
                        hl_red = True
                t = line_text.strip()
                if t:
                    lines.append({"text": t, "hl": hl_bg or hl_red,
                                  "hl_bg": hl_bg, "hl_red": hl_red})
    except:
        return []
    questions, cur = [], None
    for item in lines:
        text, hl = item["text"], item["hl"]
        m = re.match(r'^(?:C[âa]u\s*)?(\d+)[\.\:\)]\s*(.+)', text, re.IGNORECASE)
        if m and not re.match(r'^[A-Da-d][\.\)]', m.group(2)):
            if cur and len(cur["choices"]) >= 2:
                questions.append(cur)
            cur = {"id": int(m.group(1)), "question": m.group(2).strip(),
                   "choices": [], "answer": "", "explanation": ""}
        elif re.match(r'^[A-Da-d][\.\)]\s+.+', text) and cur:
            label = text[0].upper()
            cur["choices"].append({
                "label": label,
                "text": re.sub(r'^[A-Da-d][\.\)]\s*', '', text).strip(),
            })
            if hl and not cur["answer"]:
                cur["answer"] = label
        elif re.match(r'^(Đ[áa]p\s*[áa]n|ĐA)\s*[:\-]', text, re.IGNORECASE) and cur:
            mm = re.search(r'[A-D]', text.upper())
            if mm and not cur["answer"]:
                cur["answer"] = mm.group()
    if cur and len(cur["choices"]) >= 2:
        questions.append(cur)
    return questions


def parse_pdf(content: bytes) -> list:
    import fitz
    doc = fitz.open(stream=content, filetype="pdf")
    all_questions = []
    for page in doc:
        hl_rects = get_highlight_rects(page)
        if hl_rects:
            try:
                table_qs = parse_pdf_table(page, hl_rects)
                if table_qs:
                    all_questions.extend(table_qs)
                    continue
            except:
                pass
        try:
            inline_qs = parse_pdf_inline(page, hl_rects)
            all_questions.extend(inline_qs)
        except:
            pass
    for i, q in enumerate(all_questions):
        q["id"] = i + 1
    return all_questions


def _parse_choice_line(line: str) -> list:
    """Split 'A. option1\t\tB. option2\t\tC. option3\t\tD. option4' into choices."""
    import re
    choices = []
    # Match all choices on the line: A. text  B. text  C. text  D. text
    # Use tab/space as separator before each label
    parts = re.split(r'\t+', line)
    current_label = None
    current_text = ""
    for part in parts:
        part = part.strip()
        if not part:
            continue
        m = re.match(r'^([A-D])[\.\:\)]\s*(.*)', part, re.IGNORECASE)
        if m:
            if current_label:
                choices.append({"label": current_label, "text": current_text.strip()})
            current_label = m.group(1).upper()
            current_text = m.group(2)
        else:
            # Continuation of previous choice
            current_text += " " + part
    if current_label:
        choices.append({"label": current_label, "text": current_text.strip()})
    return choices


def parse_word(content: bytes) -> list:
    import docx
    import re
    doc = docx.Document(io.BytesIO(content))
    answer_map = {}

    # Parse answer tables: look for tables with "1.A", "2.B" patterns
    for table in doc.tables:
        all_text = " ".join(c.text.strip() for row in table.rows for c in row.cells)
        matches = re.findall(r'(\d+)\.?\s*([A-D])', all_text, re.IGNORECASE)
        if matches:
            for num, ans in matches:
                answer_map[int(num)] = ans.upper()
            if len(matches) > 5:
                continue

        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            for cell in cells:
                cell_matches = re.findall(r'(\d+)\.?\s*([A-D])', cell, re.IGNORECASE)
                if len(cell_matches) > 3:
                    for num, ans in cell_matches:
                        answer_map[int(num)] = ans.upper()

    questions, cur = [], None
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for i, line in enumerate(lines):
        # Check if this is a line with choices on same line (A. B. C. D.)
        choice_parts = _parse_choice_line(line)
        has_question_number = re.match(r'^(?:C[âa]u\s*)?(\d+)[\.\:\)]\s*(.+)', line, re.IGNORECASE)

        if has_question_number and choice_parts:
            # "Câu 1. Question A. opt1 B. opt2 C. opt3 D. opt4"
            q_num = int(has_question_number.group(1))
            q_text = has_question_number.group(2).strip()
            # q_text might have partial choice text, extract only the question part
            # The question text ends before the first " A. " or similar pattern
            first_choice_match = re.search(r'\s+([A-D])[\.\:\)]\s', line, re.IGNORECASE)
            if first_choice_match:
                q_text = line[:first_choice_match.start()].strip()
                # Re-parse the full line for all choices
                choice_parts = _parse_choice_line(line)

            if len(choice_parts) >= 2:
                cur = {"id": q_num, "question": q_text, "choices": choice_parts,
                       "answer": "", "explanation": ""}
                questions.append(cur)
                cur = None
            continue

        # Regular question header (on its own line)
        m = re.match(r'^(?:C[âa]u\s*)?(\d+)[\.\:\)]\s*(.+)', line, re.IGNORECASE)
        if m and not re.match(r'^[A-Da-d][\.\)]', m.group(2)):
            if cur and len(cur["choices"]) >= 2:
                questions.append(cur)
            cur = {"id": int(m.group(1)), "question": m.group(2).strip(),
                   "choices": [], "answer": "", "explanation": ""}
        elif re.match(r'^[A-Da-d][\.\)]\s+.+', line) and cur is not None:
            choices_on_line = _parse_choice_line(line)
            if choices_on_line:
                for ch in choices_on_line:
                    if ch["label"] not in [c["label"] for c in cur["choices"]]:
                        cur["choices"].append(ch)
            else:
                label = line[0].upper()
                text = re.sub(r'^[A-Da-d][\.\)]\s*', '', line).strip()
                if label not in [c["label"] for c in cur["choices"]]:
                    cur["choices"].append({"label": label, "text": text})
        elif re.match(r'^(Đ[áa]p\s*[áa]n|ĐA)', line, re.IGNORECASE) and cur is not None:
            mm = re.search(r'[A-D]', line.upper())
            if mm:
                cur["answer"] = mm.group()

    if cur and len(cur["choices"]) >= 2:
        questions.append(cur)

    for q in questions:
        if not q["answer"] and q["id"] in answer_map:
            q["answer"] = answer_map[q["id"]]

    return questions


def smart_parse(content: bytes, filename: str, api_key: str = "", force_ai: bool = False) -> dict:
    fname_lower = filename.lower()
    filetype = "pdf" if fname_lower.endswith(".pdf") else "docx"

    method, error_msg = "normal", ""
    try:
        questions = parse_pdf(content) if filetype == "pdf" else parse_word(content)
    except Exception as e:
        questions = []
        error_msg = str(e)

    total = len(questions)
    has_ans = sum(1 for q in questions if q.get("answer") in list("ABCD"))
    ans_rate = round(has_ans / total * 100) if total > 0 else 0

    ai_ok = bool(api_key and api_key != "YOUR_KEY_HERE")
    need_ai = force_ai or (ai_ok and ((total > 0 and ans_rate < 30) or total == 0))

    if need_ai and ai_ok:
        method = "vision_ai"
        try:
            questions = parse_with_vision_ai(content, filetype, api_key)
            total = len(questions)
            has_ans = sum(1 for q in questions if q.get("answer") in list("ABCD"))
            ans_rate = round(has_ans / total * 100) if total > 0 else 0
        except Exception as e:
            print(f"Vision AI error: {e}")

    return {
        "questions": questions, "method": method, "total": total,
        "has_ans": has_ans, "ans_rate": ans_rate,
        "ai_available": ai_ok, "error": error_msg,
    }


def parse_with_vision_ai(content: bytes, filetype: str, api_key: str) -> list:
    try:
        import anthropic
        from pdf2image import convert_from_bytes
    except ImportError:
        raise HTTPException(500, "Cần: pip install anthropic pdf2image")

    client = anthropic.Anthropic(api_key=api_key)
    if filetype != "pdf":
        return parse_word(content)

    try:
        images = convert_from_bytes(content, dpi=150)
    except Exception as e:
        raise HTTPException(500, f"Lỗi convert PDF→ảnh: {e}")

    VISION_PROMPT = """Đây là ảnh trang đề thi trắc nghiệm tiếng Việt.
Đáp án ĐÚNG = ô được TÔ MÀU (cyan, xanh, vàng, hồng...).
Nếu là bảng (TT | Nội dung | Đáp án A | B | C | D): cột tô màu = đáp án đúng.
Trả về JSON THUẦN TÚY không có markdown:
[{"id":1,"question":"...","choices":[{"label":"A","text":"..."},{"label":"B","text":"..."},{"label":"C","text":"..."},{"label":"D","text":"..."}],"answer":"A"}]
answer="" nếu không xác định."""

    all_qs = []
    for i, img in enumerate(images):
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=90)
        img_b64 = base64.b64encode(buf.getvalue()).decode()
        try:
            resp = client.messages.create(
                model="claude-opus-4-5", max_tokens=4096,
                messages=[{"role": "user", "content": [
                    {"type": "image", "source": {
                        "type": "base64", "media_type": "image/jpeg", "data": img_b64}},
                    {"type": "text", "text": VISION_PROMPT}
                ]}]
            )
            raw = resp.content[0].text.strip()
            m = re.search(r'\[.*\]', raw, re.DOTALL)
            if m:
                page_qs = json.loads(m.group())
                for q in page_qs:
                    q["id"] = len(all_qs) + page_qs.index(q) + 1
                all_qs.extend(page_qs)
        except Exception as e:
            print(f"Vision AI p{i+1}: {e}")
    return all_qs


# ── ENDPOINTS ──

@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    x_user_id: str = Header(default="guest"),
    x_force_ai: str = Header(default="false"),
    x_file_for: str = Header(default=""),
    current_user: Optional[dict] = Depends(get_current_user_optional),
):
    # Teacher/Admin có thể upload cho user khác
    is_privileged = current_user and current_user.get("role") in (Role.TEACHER.value, Role.ADMIN.value)
    uid = (current_user.get("sub") if current_user else None) or x_user_id
    if is_privileged and x_file_for:
        uid = x_file_for
    elif is_privileged:
        uid = current_user.get("sub")
    content = await file.read()
    fname = file.filename or "file"

    if not (fname.lower().endswith(".docx") or fname.lower().endswith(".pdf")):
        raise HTTPException(400, "Chỉ hỗ trợ .docx và .pdf")

    force_ai = x_force_ai.lower() == "true"
    from ..core.config import get_settings
    cfg = get_settings()
    ai_key = cfg.anthropic_api_key if cfg.anthropic_api_key != "YOUR_KEY_HERE" else ""

    try:
        result = smart_parse(content, fname, api_key=ai_key, force_ai=force_ai)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Lỗi xử lý: {e}")

    questions = result["questions"]
    if not questions:
        raise HTTPException(400, "Không tìm thấy câu hỏi. Thử bật AI Vision.")

    base_name = re.sub(r'\.(docx|pdf)$', '', fname, flags=re.IGNORECASE)
    file_id = re.sub(r'[^\w\-]', '_', base_name)[:50]

    index = load_json(files_index_path(uid), {})
    q_file = user_dir(uid) / f"{file_id}.json"
    existing = load_json(q_file, [])
    exist_keys = {q["question"][:60].lower() for q in existing}
    added, max_id = 0, max((q.get("id", 0) for q in existing), default=0)
    for q in questions:
        key = q["question"][:60].lower()
        if key not in exist_keys:
            max_id += 1
            q["id"] = max_id
            existing.append(q)
            exist_keys.add(key)
            added += 1

    save_json(q_file, existing)
    has_ans = sum(1 for q in existing if q.get("answer") in list("ABCD"))
    index[file_id] = {
        "name": base_name, "filename": fname,
        "count": len(existing), "with_answer": has_ans,
        "uploaded_at": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "file_id": file_id, "parse_method": result["method"],
    }
    save_json(files_index_path(uid), index)

    audit_log.log(
        EventType.FILE_UPLOAD,
        user_id=uid,
        resource=file_id,
        details={
            "filename": fname, "added": added, "total": len(existing),
            "uploaded_by": current_user.get("sub") if current_user else "guest",
        },
    )

    return {
        "file_id": file_id, "name": base_name,
        "parsed": result["total"], "added": added,
        "total_in_file": len(existing), "with_answer": has_ans,
        "ans_rate": result["ans_rate"], "parse_method": result["method"],
        "ai_available": result["ai_available"], "message": "Upload thành công",
        "for_user": uid,
    }


@router.get("/{file_id}")
def get_file(
    file_id: str,
    x_user_id: str = Header(default="guest"),
    x_file_for: str = Header(default=""),
    current_user: Optional[dict] = Depends(get_current_user_optional),
):
    uid = _resolve_file_uid(current_user, x_user_id, x_file_for)
    index = load_json(files_index_path(uid), {})
    if file_id not in index:
        raise HTTPException(404, "Không tìm thấy file")
    index[file_id]["for_user"] = uid
    return index[file_id]


@router.delete("/{file_id}")
def delete_file(
    file_id: str,
    x_user_id: str = Header(default="guest"),
    x_file_for: str = Header(default=""),
    current_user: Optional[dict] = Depends(get_current_user_optional),
):
    uid = _resolve_file_uid(current_user, x_user_id, x_file_for)
    index = load_json(files_index_path(uid), {})
    if file_id not in index:
        raise HTTPException(404, "Không tìm thấy file")
    q_file = user_dir(uid) / f"{file_id}.json"
    if q_file.exists():
        q_file.unlink()
    del index[file_id]
    save_json(files_index_path(uid), index)

    audit_log.log(EventType.FILE_DELETE, user_id=uid, resource=file_id,
                  details={"deleted_by": current_user.get("sub") if current_user else "guest"})

    return {"message": "Đã xóa"}


@router.get("/{file_id}/questions")
def get_file_questions(
    file_id: str,
    x_user_id: str = Header(default="guest"),
    x_file_for: str = Header(default=""),
    current_user: Optional[dict] = Depends(get_current_user_optional),
):
    uid = _resolve_file_uid(current_user, x_user_id, x_file_for)
    q_file = user_dir(uid) / f"{file_id}.json"
    if not q_file.exists():
        raise HTTPException(404, "Không tìm thấy file")
    questions = load_json(q_file, [])
    return {"questions": questions, "total": len(questions), "for_user": uid}


class QuestionUpdateBody(BaseModel):
    question: str
    choices: list
    answer: str = ""


@router.put("/{file_id}/questions/{q_id}")
def update_question(
    file_id: str, q_id: int, body: QuestionUpdateBody,
    x_user_id: str = Header(default="guest"),
    x_file_for: str = Header(default=""),
    current_user: Optional[dict] = Depends(get_current_user_optional),
):
    uid = _resolve_file_uid(current_user, x_user_id, x_file_for)
    q_file = user_dir(uid) / f"{file_id}.json"
    if not q_file.exists():
        raise HTTPException(404, "Không tìm thấy file")
    questions = load_json(q_file, [])
    for i, q in enumerate(questions):
        if q.get("id") == q_id:
            questions[i]["question"] = body.question
            questions[i]["choices"] = body.choices
            questions[i]["answer"] = body.answer
            save_json(q_file, questions)
            audit_log.log(EventType.QUESTION_UPDATE, user_id=uid,
                          resource=file_id, details={
                              "q_id": q_id,
                              "updated_by": current_user.get("sub") if current_user else "guest",
                          })
            return {"message": "Đã cập nhật", "question": questions[i]}
    raise HTTPException(404, "Không tìm thấy câu hỏi")


@router.delete("/{file_id}/questions/{q_id}")
def delete_question(
    file_id: str, q_id: int,
    x_user_id: str = Header(default="guest"),
    x_file_for: str = Header(default=""),
    current_user: Optional[dict] = Depends(get_current_user_optional),
):
    uid = _resolve_file_uid(current_user, x_user_id, x_file_for)
    q_file = user_dir(uid) / f"{file_id}.json"
    if not q_file.exists():
        raise HTTPException(404, "Không tìm thấy file")
    questions = load_json(q_file, [])
    new_qs = [q for q in questions if q.get("id") != q_id]
    if len(new_qs) == len(questions):
        raise HTTPException(404, "Không tìm thấy câu hỏi")
    save_json(q_file, new_qs)
    audit_log.log(EventType.QUESTION_DELETE, user_id=uid,
                  resource=file_id, details={
                      "q_id": q_id,
                      "deleted_by": current_user.get("sub") if current_user else "guest",
                  })
    return {"message": "Đã xóa câu hỏi"}


class NewQuestionBody(BaseModel):
    question: str
    choices: list
    answer: str = ""


@router.post("/{file_id}/questions")
def add_question(
    file_id: str, body: NewQuestionBody,
    x_user_id: str = Header(default="guest"),
    x_file_for: str = Header(default=""),
    current_user: Optional[dict] = Depends(get_current_user_optional),
):
    uid = _resolve_file_uid(current_user, x_user_id, x_file_for)
    q_file = user_dir(uid) / f"{file_id}.json"
    if not q_file.exists():
        raise HTTPException(404, "Không tìm thấy file")
    questions = load_json(q_file, [])
    max_id = max((q.get("id", 0) for q in questions), default=0)
    new_q = {
        "id": max_id + 1, "question": body.question,
        "choices": body.choices, "answer": body.answer,
    }
    questions.append(new_q)
    save_json(q_file, questions)
    audit_log.log(EventType.QUESTION_CREATE, user_id=uid,
                  resource=file_id, details={
                      "q_id": new_q["id"],
                      "added_by": current_user.get("sub") if current_user else "guest",
                  })
    return {"message": "Đã thêm câu hỏi", "question": new_q}
