from fastapi import FastAPI, UploadFile, File, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import json, random, re, os, base64, io, traceback, uuid, logging
from datetime import datetime, timedelta, timezone
import jwt
import importlib

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

logger = logging.getLogger(__name__)

genai = None
_GENAI_NEW = False
try:
    from google import genai as _genai_new
    genai = _genai_new
    _GENAI_NEW = True
except Exception:
    try:
        genai = importlib.import_module("google.generativeai")
    except Exception:
        genai = None

import repository as repo
from db import init_schema_from_file, get_connection

try:
    from PIL import Image
except Exception:
    Image = None

try:
    from services.science_parser import (
        detect_subject,
        normalize_chemical_formula,
        normalize_physics_units,
        enrich_science_fields,
        looks_garbled_improved,
        build_rich_text_for_display,
        classify_math_expression,
        detect_physics_elements,
        detect_chemistry_elements,
    )
    _SCIENCE_PARSER_OK = True
except ImportError:
    _SCIENCE_PARSER_OK = False
    def detect_subject(*a, **kw): return ("unknown", 0.0, {})
    def enrich_science_fields(q): return q
    def looks_garbled_improved(t): return False

try:
    from services.embedding_service import EmbeddingService
    from services.vector_store import VectorStore
    from services.rag_service import RAGService
    from services.mrc_service import MRCService
    from services.pipeline_service import AIPipeline

    _AI_MODULES_IMPORTED = True
except ImportError as e:
    print(f"[WARN] AI services chưa sẵn sàng: {e}")
    _AI_MODULES_IMPORTED = False

def _get_pix2tex_model():
    return False

def _formula_ocr_from_pixmap(pm) -> str:
    return ""

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()

app = FastAPI()
_cors_origins = [o.strip() for o in os.getenv("CORS_ORIGINS", "*").split(",") if o.strip()]
if not _cors_origins:
    _cors_origins = ["*"]
app.add_middleware(CORSMiddleware, allow_origins=_cors_origins, allow_methods=["*"], allow_headers=["*"])

JWT_SECRET = os.getenv("JWT_SECRET", "change-me-in-production")
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = int(os.getenv("JWT_EXPIRE_HOURS", "24"))
JWT_REFRESH_SECRET = os.getenv("JWT_REFRESH_SECRET", JWT_SECRET)
JWT_REFRESH_EXPIRE_DAYS = int(os.getenv("JWT_REFRESH_EXPIRE_DAYS", "14"))

_embedding_service = None
_vector_store = None
_rag_service = None
_mrc_service = None
_ai_pipeline = None

def _init_ai_services():
    """Khởi tạo AI services với khởi tạo song song (parallel init)."""
    global _embedding_service, _vector_store, _rag_service, _mrc_service, _ai_pipeline

    if _ai_pipeline is not None:
        return

    try:
        from services.embedding_service import EmbeddingService
        from services.vector_store import VectorStore
        from services.rag_service import RAGService
        from services.mrc_service import MRCService
        from services.pipeline_service import AIPipeline

        print("[CobraQ] AI Pipeline: Đang khởi tạo (lần đầu, cần tải model...)...")
        print("[CobraQ]   - Dang tai embedding model (neu chua co cache)...")

        _embedding_service = EmbeddingService(
            model_name=os.getenv("EMBEDDING_MODEL", "paraphrase-multilingual-MiniLM-L12-v2")
        )

        print("[CobraQ]   - Embedding model da san sang!")
        print("[CobraQ]   - Khoi tao VectorStore & RAG...")

        _vector_store = VectorStore(
            persist_directory=os.getenv("VECTOR_DB_PATH", "./chroma_db")
        )

        _rag_service = RAGService(
            vector_store=_vector_store,
            embedding_service=_embedding_service,
            similarity_threshold=float(os.getenv("RAG_SIMILARITY_THRESHOLD", "0.85"))
        )

        print("[CobraQ]   - VectorStore & RAG da san sang!")
        print("[CobraQ]   - Khoi tao Groq MRC...")

        groq_key = GROQ_API_KEY
        if groq_key:
            _mrc_service = MRCService(api_key=groq_key)
        else:
            print("[WARN] Không có GROQ_API_KEY hợp lệ, AI pipeline sẽ không hoạt động")
            return

        _ai_pipeline = AIPipeline(
            embedding_service=_embedding_service,
            vector_store=_vector_store,
            rag_service=_rag_service,
            mrc_service=_mrc_service,
            cache_llm_results=True,
            default_threshold=0.90
        )

        print("[CobraQ] AI Pipeline da san sang (RAG + MRC)")

        _warmup_embedding()
    except Exception as e:
        import traceback
        print(f"[ERROR] Không thể khởi tạo AI services: {e}")
        traceback.print_exc()


def _warmup_embedding():
    """Warmup embedding model bang 1 request nho de trigger JIT/caching."""
    import threading
    def _do_warmup():
        try:
            sample = _embedding_service.embed_text("Khoi dong nhanh")
        except Exception:
            pass
    t = threading.Thread(target=_do_warmup, daemon=True)
    t.start()


def _is_ai_pipeline_ready():
    """Check if AI pipeline is initialized and ready."""
    return _ai_pipeline is not None


def _build_token(user: dict, token_type: str, expires_delta: timedelta, secret: str) -> str:
    now = datetime.now(timezone.utc)
    payload = {
        "sub": user["uid"],
        "email": user.get("email"),
        "role": user.get("role") or "user",
        "type": token_type,
        "iat": int(now.timestamp()),
        "exp": int((now + expires_delta).timestamp()),
    }
    return jwt.encode(payload, secret, algorithm=JWT_ALGORITHM)


def create_access_token(user: dict) -> str:
    return _build_token(user, "access", timedelta(hours=JWT_EXPIRE_HOURS), JWT_SECRET)


def create_refresh_token(user: dict) -> str:
    return _build_token(user, "refresh", timedelta(days=JWT_REFRESH_EXPIRE_DAYS), JWT_REFRESH_SECRET)


def decode_token(token: str, secret: str, expected_type: str) -> dict:
    try:
        payload = jwt.decode(token, secret, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(401, "Token đã hết hạn")
    except jwt.InvalidTokenError:
        raise HTTPException(401, "Token không hợp lệ")
    if (payload.get("type") or "") != expected_type:
        raise HTTPException(401, "Sai loại token")
    if repo.is_token_revoked(token):
        raise HTTPException(401, "Token đã bị thu hồi")
    return payload


def decode_access_token(token: str) -> dict:
    return decode_token(token, JWT_SECRET, "access")


def decode_refresh_token(token: str) -> dict:
    return decode_token(token, JWT_REFRESH_SECRET, "refresh")


def get_auth_user(authorization: str = Header(default="")) -> dict:
    auth = (authorization or "").strip()
    if not auth.lower().startswith("bearer "):
        raise HTTPException(401, "Thiếu Bearer token")
    token = auth[7:].strip()
    if not token:
        raise HTTPException(401, "Thiếu Bearer token")
    payload = decode_access_token(token)
    uid = (payload.get("sub") or "").strip()
    if not uid:
        raise HTTPException(401, "Token không hợp lệ")
    user = repo.get_user_by_uid(uid)
    if not user:
        raise HTTPException(401, "Người dùng không tồn tại")
    return user


@app.on_event("startup")
def _startup_init_db():
    init_schema_from_file()
    repo.cleanup_revoked_tokens()

    try:
        from db import _engine as _db_engine
        eng = _db_engine()
    except Exception:
        eng = "unknown"
    print(f"[CobraQ] DB engine: {eng}")
    if has_valid_key():
        print("[CobraQ] Groq AI: ON (đã có GROQ_API_KEY)")
        print("[CobraQ] AI Pipeline: Dang tai model... (chay ngam, server van tra loi ngay)")
        import threading
        t = threading.Thread(target=_init_ai_services, daemon=True)
        t.start()
    else:
        print("[CobraQ] Groq AI: OFF (chưa có GROQ_API_KEY hợp lệ)")

    app_env = (os.getenv("APP_ENV", "dev") or "dev").strip().lower()
    if app_env in ("prod", "production") and JWT_SECRET == "change-me-in-production":
        raise RuntimeError("JWT_SECRET không được để mặc định ở môi trường production")

    admin_email = (os.getenv("ADMIN_EMAIL", "") or "").strip().lower()
    admin_password = (os.getenv("ADMIN_PASSWORD", "") or "").strip()
    if admin_email and admin_password:
        existing = repo.get_user_by_email(admin_email)
        if existing:
            if (existing.get("role") or "user") != "admin":
                repo.set_user_role(existing["uid"], "admin")
        else:
            uid = f"adm_{uuid.uuid4().hex[:20]}"
            try:
                repo.register_user(uid, admin_email, admin_password, role="admin")
            except Exception as e:
                print(f"Seed admin thất bại: {e}")


def load_config():
    cfg = {"groq_key": GROQ_API_KEY, "ai_parse_enabled": True, "ai_fill_enabled": False}
    try:
        cfg["ai_parse_enabled"] = repo.get_ai_parse_enabled()
        cfg["ai_fill_enabled"] = bool(repo.get_config_value("ai_fill_enabled", 0))
    except Exception:
        pass
    return cfg

def has_valid_key():
    k = GROQ_API_KEY.strip()
    return bool(k and k != "ĐIỀN_API_KEY_CỦA_BẠN_VÀO_ĐÂY" and k != "YOUR_KEY_HERE" and "..." not in k)

def is_red_text(color) -> bool:
    if color is None:
        return False
    try:
        if isinstance(color, (list, tuple)) and len(color) >= 3:
            r, g, b = float(color[0]), float(color[1]), float(color[2])
            return r > 0.65 and g < 0.4 and b < 0.4
        if isinstance(color, int) and color > 0:
            r = ((color >> 16) & 0xFF) / 255.0
            g = ((color >> 8) & 0xFF) / 255.0
            b = (color & 0xFF) / 255.0
            return r > 0.65 and g < 0.4 and b < 0.4
    except Exception:
        pass
    return False


CHOICE_INLINE_PAT = re.compile(
    r"(?:^|(?<=\s)|(?<=[\.．,;:]))\b([A-D])[\.\)．:：]\s*",
)


def _sanitize_choice_text(t: str) -> str:
    if not t:
        return ""
    t = str(t).replace("\u00a0", " ").replace("\u202f", " ")
    t = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", t)
    return t.strip()


def split_choice_segments(text: str) -> list:
    if not text or not str(text).strip():
        return []
    s = _sanitize_choice_text(text)
    matches = list(CHOICE_INLINE_PAT.finditer(s))
    if not matches:
        return []
    out = []
    for i, m in enumerate(matches):
        label = m.group(1).upper()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(s)
        chunk = s[start:end].strip()
        if chunk:
            out.append({"label": label, "text": chunk})
    return out


def _expand_choice_cell(text: str, cell_label: str) -> list:
    """Một ô đáp án (A/B/C/D) có thể chứa cả mốc D. E. — phần trước mốc đầu thuộc cell_label."""
    t = _sanitize_choice_text(text or "")
    cl = (cell_label or "?").strip().upper()
    if not t:
        return []
    matches = list(CHOICE_INLINE_PAT.finditer(t))
    if not matches:
        return [{"label": cl, "text": t}] if cl in "ABCD" else []
    out = []
    if matches[0].start() > 0 and cl in "ABCD":
        pre = t[: matches[0].start()].strip()
        if pre:
            out.append({"label": cl, "text": pre})
    for i, m in enumerate(matches):
        lb = m.group(1).upper()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(t)
        chunk = t[start:end].strip()
        if chunk:
            out.append({"label": lb, "text": chunk})
    if not out:
        return [{"label": cl, "text": t}] if cl in "ABCD" else []
    return out


def normalize_merged_choices_in_question(q: dict) -> None:
    chs = q.get("choices") or []
    if not chs:
        return
    for _ in range(5):
        new_chs = []
        seen = set()
        for ch in chs:
            label = (ch.get("label") or "?").strip().upper()
            t = _sanitize_choice_text(ch.get("text") or "")
            for p in _expand_choice_cell(t, label):
                lb = (p.get("label") or "?").strip().upper()
                tx = _sanitize_choice_text(p.get("text") or "")
                tx = re.sub(r"^\s*(?:Phương\s*án|Đáp\s*án)\s*[A-D]\s*[:\-\.]\s*", "", tx, flags=re.IGNORECASE)
                tx = re.sub(r"\s+", " ", tx).strip(" .;:\t\n\r")
                if lb in "ABCD" and lb not in seen and tx:
                    seen.add(lb)
                    new_chs.append({"label": lb, "text": tx})
        if new_chs:
            new_chs = sorted(new_chs, key=lambda x: "ABCD".index(x["label"]) if x["label"] in "ABCD" else 9)
        if not new_chs:
            break
        same_len = len(new_chs) == len(chs)
        same_txt = same_len and all(
            (chs[i].get("label"), _sanitize_choice_text(chs[i].get("text") or ""))
            == (new_chs[i]["label"], new_chs[i]["text"])
            for i in range(len(new_chs))
        )
        q["choices"] = new_chs
        if same_txt:
            break
        chs = new_chs


def _clean_choices_payload(choices: list) -> list:
    out = []
    seen = set()
    for c in choices or []:
        if not isinstance(c, dict):
            continue
        lb = (c.get("label") or "").strip().upper()
        tx = _sanitize_choice_text(c.get("text") or "")
        if lb in "ABCD" and tx and lb not in seen:
            seen.add(lb)
            out.append({"label": lb, "text": tx})
    return sorted(out, key=lambda x: "ABCD".index(x["label"]))


def _looks_like_math_expr(s: str) -> bool:
    t = (s or "").strip()
    if not t:
        return False
    math_tokens = [r"\\frac", r"\\sqrt", r"\\int", r"\\sum", r"\\lim", r"\\alpha", r"\\beta", r"\\theta", r"\\pi", r"\^", r"_", r"=", r"≤", r"≥", r"∫", r"Σ", r"√", r"→", r"↔", r"∀", r"∃", r"\bcos\b", r"\bsin\b", r"\btan\b", r"\blog\b", r"\bln\b"]
    return any(re.search(p, t, flags=re.IGNORECASE) for p in math_tokens)


def _to_rich_inline(text: str) -> str:
    t = (text or "").strip()
    if not t:
        return ""
    if "$" in t or "\\(" in t or "\\[" in t:
        return t
    if _looks_like_math_expr(t):
        return f"\\({t}\\)"
    return t


def enrich_question_rich_fields(q: dict) -> dict:
    qq = dict(q or {})
    question = str(qq.get("question") or "").strip()
    choices = list(qq.get("choices") or [])

    question_rich_raw = str(qq.get("question_rich") or "").strip()
    question_rich = question_rich_raw or _to_rich_inline(question)

    choices_rich = []
    math_hits = 1 if (_looks_like_math_expr(question) or _looks_like_math_expr(question_rich)) else 0
    for c in choices:
        txt = str((c or {}).get("text") or "").strip()
        txt_rich_raw = str((c or {}).get("text_rich") or "").strip()
        txt_rich = txt_rich_raw or _to_rich_inline(txt)
        if _looks_like_math_expr(txt) or _looks_like_math_expr(txt_rich):
            math_hits += 1
        choices_rich.append({"label": (c.get("label") or "").strip().upper(), "text": txt_rich})

    confidence = 0.82
    flags = {"math_detected": bool(math_hits), "source": "cpu_only_v1"}
    if math_hits >= 2:
        confidence = 0.9
    if not question or len(choices) < 2:
        confidence = 0.55
        flags["needs_review"] = True

    qq["question_rich"] = question_rich or question
    qq["choices_rich"] = choices_rich
    qq["parse_confidence"] = round(confidence, 4)
    qq["parse_flags"] = flags

    if _SCIENCE_PARSER_OK:
        qq = enrich_science_fields(qq)

    return qq


def _validate_question_choices(choices: list, answer: str) -> tuple[list, str]:
    clean = _clean_choices_payload(choices)
    if len(clean) < 2:
        raise HTTPException(400, "Cần ít nhất 2 phương án có nội dung (C, D có thể để trống).")
    ans = (answer or "").strip().upper()
    if ans:
        if ans not in list("ABCD") or ans not in {c["label"] for c in clean}:
            raise HTTPException(400, "Đáp án đúng phải là A/B/C/D và trùng một phương án đã nhập.")
    return clean, ans



def is_highlight(fill):
    if not fill or len(fill) < 3: return False
    r, g, b = float(fill[0]), float(fill[1]), float(fill[2])
    if r > 0.92 and g > 0.92 and b > 0.92: return False  
    if r < 0.05 and g < 0.05 and b < 0.05: return False  
    if r > 0.85 and g > 0.85 and b > 0.85: return False  
    if g > 0.50 and b > 0.40 and r < 0.60: return True
    if g > 0.55 and g > r * 1.3 and g > b * 1.2: return True
    if r > 0.75 and g > 0.70 and b < 0.40: return True
    if r > 0.80 and g < 0.65 and b < 0.65: return True
    if b > 0.55 and b > r * 1.2 and g > 0.45: return True
    return False

def get_highlight_rects(page):
    rects = []
    try:
        for d in page.get_drawings():
            if is_highlight(d.get("fill")):
                rects.append(d["rect"])
    except: pass
    return rects

def span_in_rect(bbox, rects, pad=4):
    cx = (bbox[0] + bbox[2]) / 2
    cy = (bbox[1] + bbox[3]) / 2
    for r in rects:
        if r.x0-pad <= cx <= r.x1+pad and r.y0-pad <= cy <= r.y1+pad:
            return True
    return False

def get_all_spans(page):
    spans = []
    try:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0: continue
            for line in block["lines"]:
                for sp in line["spans"]:
                    t = sp["text"].strip()
                    if t:
                        bbox = sp["bbox"]
                        spans.append({"text": t, "bbox": bbox, "cx": (bbox[0]+bbox[2])/2, "cy": (bbox[1]+bbox[3])/2})
    except: pass
    return spans

def group_rows(spans, tol=12):
    rows = []
    for sp in sorted(spans, key=lambda s: s["cy"]):
        placed = False
        for row in rows:
            if abs(row["cy"] - sp["cy"]) < tol:
                row["spans"].append(sp)
                row["cy"] = (row["cy"] + sp["cy"]) / 2
                placed = True
                break
        if not placed:
            rows.append({"cy": sp["cy"], "spans": [sp]})
    return sorted(rows, key=lambda r: r["cy"])

def parse_pdf_table(page, hl_rects):
    spans = get_all_spans(page)
    if not spans: return None
    rows = group_rows(spans)
    col_A = col_B = col_C = col_D = col_Q = None
    header_y = None

    for row in rows:
        has_da = bool(re.search(r'ĐÁPÁN|ANSWER', " ".join(s["text"] for s in row["spans"]), re.IGNORECASE))
        labels = [s for s in row["spans"] if s["text"].strip().upper() in ["A", "B", "C", "D"]]
        if has_da or len(labels) >= 3:
            header_y = row["cy"]
            for sp in row["spans"]:
                t = sp["text"].upper().replace(" ", "")
                if any(k in t for k in ["NỘIDUNG", "CÂUHỎI", "CONTENT", "NOIUNG"]): col_Q = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]NA|ANSWERA', t) or t == "A": col_A = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]NB|ANSWERB', t) or t == "B": col_B = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]NC|ANSWERC', t) or t == "C": col_C = sp["cx"]
                if re.search(r'[ĐD][AÁ]P[AÁ]ND|ANSWERD', t) or t == "D": col_D = sp["cx"]
            break

    if sum(1 for c in [col_A, col_B, col_C, col_D] if c is not None) < 2: return None

    col_map = {k: v for k, v in [("A", col_A), ("B", col_B), ("C", col_C), ("D", col_D)] if v is not None}
    answer_map = {}
    for hl in hl_rects:
        hcx, hcy = (hl.x0 + hl.x1) / 2, (hl.y0 + hl.y1) / 2
        if header_y and hcy <= header_y + 5: continue
        best_label, best_dist = None, 999
        for label, cx in col_map.items():
            if abs(hcx - cx) < best_dist:
                best_dist, best_label = abs(hcx - cx), label
        if not best_label or best_dist > 120: continue

        best_tt, best_dy = None, 999
        for sp in spans:
            if re.match(r'^\d+$', sp["text"]):
                if abs(sp["cy"] - hcy) < best_dy:
                    best_dy, best_tt = abs(sp["cy"] - hcy), int(sp["text"])
        if best_tt and best_dy < 30: answer_map[best_tt] = best_label

    if not answer_map: return None

    questions, data_rows = [], [r for r in rows if header_y is None or r["cy"] > header_y + 8]
    tt_to_spans, current_tt = {}, None
    for row in data_rows:
        tt_spans = [s for s in row["spans"] if re.match(r'^\d+$', s["text"])]
        if tt_spans: current_tt = int(tt_spans[0]["text"])
        if current_tt: tt_to_spans.setdefault(current_tt, []).extend(row["spans"])

    for tt, row_spans in tt_to_spans.items():
        if tt not in answer_map: continue
        if col_Q:
            q_spans = [s for s in row_spans if col_Q - 180 <= s["cx"] <= col_Q + 180 and s["text"] not in [str(tt)]]
        else:
            q_spans = [s for s in row_spans if len(s["text"]) > 2 and not re.match(r'^[\dABCD]$', s["text"]) and not any(abs(s["cx"] - cx) < 80 for cx in col_map.values())]

        q_text = " ".join(s["text"] for s in sorted(q_spans, key=lambda x: (x["cy"], x["cx"]))).strip()
        if not q_text: continue

        choices = []
        for label in sorted(col_map.keys()):
            cx = col_map[label]
            c_spans = [s for s in row_spans if abs(s["cx"] - cx) < 90 and s["text"] not in [str(tt)]]
            c_text = " ".join(s["text"] for s in sorted(c_spans, key=lambda x: (x["cy"], x["cx"]))).strip()
            if c_text: choices.append({"label": label, "text": c_text})

        if len(choices) >= 2:
            questions.append({"id": tt, "question": q_text, "choices": choices, "answer": answer_map.get(tt, ""), "explanation": ""})
    return questions if questions else None

def parse_pdf_inline(page, hl_rects):
    lines = []
    try:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                line_text, hl_bg, hl_red = "", False, False
                x0, y0, x1, y1 = 10**9, 10**9, -1, -1
                for sp in line["spans"]:
                    line_text += sp["text"]
                    bx = sp.get("bbox") or [0, 0, 0, 0]
                    x0, y0 = min(x0, bx[0]), min(y0, bx[1])
                    x1, y1 = max(x1, bx[2]), max(y1, bx[3])
                    if hl_rects and span_in_rect(sp["bbox"], hl_rects):
                        hl_bg = True
                    if is_red_text(sp.get("color", 0)):
                        hl_red = True
                t = line_text.strip()
                if not t:
                    continue
                lines.append({"text": t, "rich": t, "hl": hl_bg or hl_red})
    except:
        return []

    questions, cur = [], None
    for item in lines:
        text, hl = item["text"], item["hl"]
        rich_text = item.get("rich") or text
        m = re.match(r"^(?:C[âa]u\s*)(\d+)[\.\:\)]\s*(.+)", text, re.IGNORECASE)
        m_loose = re.match(r"^(\d+)[\.\:\)]\s*(.+)", text, re.IGNORECASE)
        can_start_loose = False
        if m_loose:
            probe_rest = (m_loose.group(2) or "").strip()
            can_start_loose = len(split_choice_segments(probe_rest)) >= 3

        mm = m if m else (m_loose if can_start_loose else None)
        if mm and not re.match(r"^[A-Da-d][\.\)]", mm.group(2)):
            if cur and len(cur["choices"]) >= 2:
                questions.append(cur)
            rest = mm.group(2).strip()
            embedded = split_choice_segments(rest)
            if len(embedded) >= 2:
                fm = CHOICE_INLINE_PAT.search(rest)
                q_stem = rest[: fm.start()].strip() if fm else rest
                cur = {"id": int(mm.group(1)), "question": q_stem, "question_rich": rich_text, "choices": list(embedded), "answer": "", "explanation": ""}
            else:
                cur = {"id": int(mm.group(1)), "question": rest, "question_rich": rich_text, "choices": [], "answer": "", "explanation": ""}
        elif cur:
            parts = split_choice_segments(text)
            if parts:
                for p in parts:
                    p2 = dict(p)
                    p2["text_rich"] = p2.get("text") or ""
                    cur["choices"].append(p2)
                    if hl and not cur["answer"]:
                        cur["answer"] = p2["label"]
            elif re.match(r"^[A-Da-d][\.\)]\s*", text):
                label = text[0].upper()
                choice_text = re.sub(r"^[A-Da-d][\.\)]\s*", "", text).strip()
                cur["choices"].append({"label": label, "text": choice_text, "text_rich": (rich_text if _looks_like_math_expr(choice_text) else choice_text)})
                if hl and not cur["answer"]:
                    cur["answer"] = label
            elif re.match(r"^(Đ[áa]p\s*[áa]n|ĐA)\s*[:\-]", text, re.IGNORECASE):
                mm = re.search(r"[A-D]", text.upper())
                if mm and not cur["answer"]:
                    cur["answer"] = mm.group()
    if cur and len(cur["choices"]) >= 2:
        questions.append(cur)
    for q in questions:
        normalize_merged_choices_in_question(q)
    return questions

def _looks_like_math_question_text(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return False
    if re.search(r"\d", t):
        return True
    math_kw = [
        "hàm số", "đạo hàm", "tích phân", "phương trình", "bất phương trình", "log", "sin", "cos", "tan",
        "đường thẳng", "mặt phẳng", "thể tích", "xác suất", "tọa độ", "cực trị", "nghiệm", "giới hạn",
        "số phức", "ma trận", "vectơ", "vector", "cấp số", "hình học", "tham số", "biến số",
    ]
    if any(k in t for k in math_kw):
        return True
    if re.search(r"[xyztmna-b]\b|[=<>≤≥±√∫Σπ]", t):
        return True
    return False


def _is_likely_valid_mcq(q: dict) -> bool:
    if not isinstance(q, dict):
        return False
    question = str(q.get("question") or "").strip()
    choices = list(q.get("choices") or [])

    if len(question) < 10 or len(question) > 1200:
        return False
    if re.search(r"https?://|www\.|@", question, re.IGNORECASE):
        return False

    if len(choices) < 3:
        return False

    if re.search(r"thời\s+gian\s+tập\s+nhảy|chỉ\s+rất\s+thích\s+nhảy|chất\s+lượng\s+sữa|khảo\s+sát|đoạn\s+văn", question, re.IGNORECASE):
        return False

    labels = []
    non_empty = 0
    for c in choices:
        lb = str((c or {}).get("label") or "").strip().upper()
        tx = str((c or {}).get("text") or "").strip()
        if lb:
            labels.append(lb)
        if tx and len(tx) >= 2:
            non_empty += 1

    if non_empty < 3:
        return False
    if any(lb not in ("A", "B", "C", "D") for lb in labels if lb):
        return False

    return True


def _looks_garbled_text(s: str) -> bool:
    t = (s or "").strip()
    if not t:
        return False
    bad_patterns = [r"[%#@\^~]{2,}", r"[A-Za-z0-9]{1,2}[%#@][A-Za-z0-9]{1,2}", r"K%|Ã|Â|ð|�"]
    if any(re.search(p, t) for p in bad_patterns):
        return True
    weird = re.findall(r"[^\w\s\+\-\*/=\(\)\[\]\{\}\.,:;!?<>&%√πΣ∫≤≥±°'\"\u00C0-\u1EF9]", t, flags=re.UNICODE)
    return len(weird) >= 3


def _garbled_ratio(questions: list) -> float:
    if not questions:
        return 1.0
    bad = 0
    total = 0
    for q in questions[:120]:
        total += 1
        qq = str(q.get("question") or "")
        garbled = _looks_garbled_text(qq)
        if _SCIENCE_PARSER_OK and not garbled:
            garbled = looks_garbled_improved(qq)
        if garbled:
            bad += 1
            continue
        for c in (q.get("choices") or []):
            ct = str((c or {}).get("text") or "")
            garbled_c = _looks_garbled_text(ct)
            if _SCIENCE_PARSER_OK and not garbled_c:
                garbled_c = looks_garbled_improved(ct)
            if garbled_c:
                bad += 1
                break
    return (bad / total) if total else 1.0


def parse_pdf(content: bytes) -> list:
    import fitz
    doc = fitz.open(stream=content, filetype="pdf")
    all_questions = []
    for page in doc:
        hl_rects = get_highlight_rects(page)
        if hl_rects:
            try:
                table_qs = parse_pdf_table(page, hl_rects)
                if table_qs:
                    all_questions.extend(table_qs)
                    continue
            except:
                pass
        try:
            all_questions.extend(parse_pdf_inline(page, hl_rects))
        except:
            pass

    cleaned = []
    for q in all_questions:
        normalize_merged_choices_in_question(q)
        if _is_likely_valid_mcq(q):
            cleaned.append(q)

    dedup = []
    seen = set()
    for q in cleaned:
        key = re.sub(r"\s+", " ", str(q.get("question") or "").strip().lower())
        if not key or key in seen:
            continue
        seen.add(key)
        dedup.append(q)

    for i, q in enumerate(dedup):
        q["id"] = i + 1

    return dedup


def _parse_lines(lines):
    questions, cur = [], None
    for line in lines:
        line = line.strip() if isinstance(line, str) else ""
        if not line:
            continue
        m = re.match(r"^(?:C[âa]u\s*)?(\d+)[\.\:\)]\s*(.+)", line, re.IGNORECASE)
        if m and not re.match(r"^[A-Da-d][\.\)]", m.group(2)):
            if cur and len(cur["choices"]) >= 2:
                questions.append(cur)
            rest = m.group(2).strip()
            embedded = split_choice_segments(rest)
            if len(embedded) >= 2:
                fm = CHOICE_INLINE_PAT.search(rest)
                q_stem = rest[: fm.start()].strip() if fm else rest
                cur = {"id": int(m.group(1)), "question": q_stem, "choices": list(embedded), "answer": "", "explanation": ""}
            else:
                cur = {"id": int(m.group(1)), "question": rest, "choices": [], "answer": "", "explanation": ""}
        elif cur:
            parts = split_choice_segments(line)
            if parts:
                for p in parts:
                    cur["choices"].append(p)
            elif re.match(r"^[A-Da-d][\.\)]\s*", line):
                cur["choices"].append({"label": line[0].upper(), "text": re.sub(r"^[A-Da-d][\.\)]\s*", "", line).strip()})
            elif re.match(r"^(Đ[áa]p\s*[áa]n|ĐA)", line, re.IGNORECASE):
                mm = re.search(r"[A-D]", line.upper())
                if mm:
                    cur["answer"] = mm.group()
    if cur and len(cur["choices"]) >= 2:
        questions.append(cur)
    for q in questions:
        normalize_merged_choices_in_question(q)
    return questions


def _extract_answer_map_from_lines(lines: list[str]) -> dict:
    """Trích bảng đáp án từ text thường: fuzzy cho nhiều format khác nhau."""
    ans_map = {}

    def _normalize_line(x: str) -> str:
        s = (x or "").replace("\u00a0", " ").replace("\u202f", " ")
        s = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", s)
        s = re.sub(r"[\|,;]+", " ", s)
        s = re.sub(r"\s+", " ", s).strip()
        return s

    for ln in lines:
        s = _normalize_line((ln or "").upper())
        if not s:
            continue

        for n, a in re.findall(r"\b(?:CÂU\s*)?(\d{1,4})\s*(?:[\)\].:\-=>]|\s)*(?:ĐÁP\s*ÁN\s*)?([ABCD])\b", s):
            ans_map[int(n)] = a

        for n, a in re.findall(r"\bCÂU\s*(\d{1,4})\b[^A-D\n]{0,25}\b([ABCD])\b", s):
            ans_map[int(n)] = a

    for i in range(len(lines) - 1):
        n_line = _normalize_line((lines[i] or ""))
        a_line = _normalize_line((lines[i + 1] or "").upper())
        nums = re.findall(r"\b\d{1,4}\b", n_line)
        letters = re.findall(r"\b[ABCD]\b", a_line)
        if nums and letters and len(nums) == len(letters):
            for n, a in zip(nums, letters):
                ans_map[int(n)] = a

    return ans_map


def parse_docx_without_docx(content: bytes) -> list:
    import html
    import io
    import re
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception:
        return []

    def _cell_text(cell_xml: str) -> str:
        s = re.sub(r"</w:p>", "\n", cell_xml)
        s = re.sub(r"<[^>]+>", "", s)
        s = html.unescape(s)
        return " ".join([x.strip() for x in s.splitlines() if x.strip()]).strip()

    def _cell_marked(cell_xml: str) -> bool:
        if re.search(r"<w:b(?:\s|/|>)", cell_xml):
            return True
        m = re.search(r"<w:color[^>]*w:val=\"([0-9A-Fa-f]{6}|auto)\"", cell_xml)
        if m and m.group(1).lower() not in ("000000", "auto"):
            return True
        if re.search(r"<w:highlight\b", cell_xml):
            return True
        return False

    tr_blocks = re.findall(r"<w:tr\b[\s\S]*?</w:tr>", xml)
    questions = []
    qid = 0
    for tr in tr_blocks:
        tc_blocks = re.findall(r"<w:tc\b[\s\S]*?</w:tc>", tr)
        if len(tc_blocks) < 6:
            continue

        vals = [_cell_text(tc) for tc in tc_blocks]
        first = (vals[0] or "").strip().upper()
        if first in ("TT", "STT", "SỐ TT") or any("NỘI DUNG CÂU HỎI" in (v or "").upper() for v in vals):
            continue
        if not re.match(r"^\d+$", vals[0] or ""):
            continue

        q_text = (vals[2] if len(vals) > 2 else "").strip()
        a_txt = (vals[3] if len(vals) > 3 else "").strip()
        b_txt = (vals[4] if len(vals) > 4 else "").strip()
        c_txt = (vals[5] if len(vals) > 5 else "").strip()
        d_txt = (vals[6] if len(vals) > 6 else "").strip()

        if not q_text or len(q_text) < 6:
            continue

        labels = ["A", "B", "C", "D"]
        texts = [a_txt, b_txt, c_txt, d_txt]
        choices = [{"label": labels[i], "text": texts[i]} for i in range(4) if texts[i]]
        if len(choices) < 2:
            continue

        marked = []
        for i, idx in enumerate((3, 4, 5, 6)):
            if idx < len(tc_blocks) and texts[i] and _cell_marked(tc_blocks[idx]):
                marked.append(labels[i])
        answer = marked[0] if marked else ""

        qid += 1
        questions.append({"id": qid, "question": q_text, "choices": choices, "answer": answer, "explanation": ""})

    x = re.sub(r"</w:p>", "\n", xml)
    x = re.sub(r"</w:tr>", "\n", x)
    x = re.sub(r"</w:tc>", "\n", x)
    text = re.sub(r"<[^>]+>", "", x)
    text = html.unescape(text)
    lines = [ln.strip() for ln in text.replace("\r", "\n").split("\n") if ln.strip()]

    if questions:
        answer_map = _extract_answer_map_from_lines(lines)
        for q in questions:
            if not q.get("answer") and q.get("id") in answer_map:
                q["answer"] = answer_map[q["id"]]
            normalize_merged_choices_in_question(q)
        return questions

    parsed = _parse_lines(lines)
    answer_map = _extract_answer_map_from_lines(lines)
    for q in parsed:
        if not q.get("answer") and q.get("id") in answer_map:
            q["answer"] = answer_map[q["id"]]
    return parsed


def parse_word(content: bytes) -> list:
    import io
    import importlib
    try:
        docx = importlib.import_module("docx")
    except Exception:
        return parse_docx_without_docx(content)

    doc = docx.Document(io.BytesIO(content))
    questions = []
    max_id = 0
    current_q = None

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 7:
                cell_lines = []
                for c in cells:
                    lines = [x.strip() for x in c.text.split('\n') if x.strip()]
                    cell_lines.append(lines)
                
                correct_keywords = [[], [], [], []]
                correct_marked = [False, False, False, False]
                for j in range(4):
                    idx = 3 + j
                    if idx < len(cells):
                        for p in cells[idx].paragraphs:
                            for run in p.runs:
                                rt = run.text.strip()
                                marked = bool(
                                    run.bold
                                    or (run.font.color and run.font.color.rgb and str(run.font.color.rgb) not in ("000000", "None"))
                                    or (run.font.highlight_color is not None)
                                )
                                if marked:
                                    correct_marked[j] = True
                                if rt and marked and len(rt) > 0:
                                    correct_keywords[j].append(rt)
                
                max_lines = max(len(lines) for lines in cell_lines) if cell_lines else 0
                for i in range(max_lines):
                    tt_text = cell_lines[0][i] if i < len(cell_lines[0]) else ""
                    
                    if tt_text.strip().upper() in ["TT", "STT", "SỐ TT"]:
                        continue

                    if tt_text.isdigit():
                        if current_q and current_q["choices"]:
                            questions.append(current_q)
                            max_id = max(max_id, current_q["id"])
                        
                        labels = ['A', 'B', 'C', 'D']
                        choices = []
                        correct_ans = ""
                        for j in range(4):
                            idx = 3 + j
                            ans_text = cell_lines[idx][i] if idx < len(cell_lines) and i < len(cell_lines[idx]) else ""
                            if ans_text:
                                choices.append({"label": labels[j], "text": ans_text})
                                if not correct_ans and correct_marked[j]:
                                    correct_ans = labels[j]
                                if not correct_ans and any(k in ans_text for k in correct_keywords[j]):
                                    correct_ans = labels[j]
                        
                        q_text = cell_lines[2][i] if 2 < len(cell_lines) and i < len(cell_lines[2]) else ""
                        current_q = {
                            "id": int(tt_text),
                            "question": q_text,
                            "choices": choices,
                            "answer": correct_ans,
                            "explanation": ""
                        }
                    else:
                        if current_q:
                            q_text = cell_lines[2][i] if 2 < len(cell_lines) and i < len(cell_lines[2]) else ""
                            if q_text: current_q["question"] += "\n" + q_text
                            
                            for j in range(4):
                                idx = 3 + j
                                ans_text = cell_lines[idx][i] if idx < len(cell_lines) and i < len(cell_lines[idx]) else ""
                                if ans_text and j < len(current_q["choices"]):
                                    current_q["choices"][j]["text"] += "\n" + ans_text
                                    if not current_q["answer"] and correct_marked[j]:
                                        current_q["answer"] = ['A', 'B', 'C', 'D'][j]
                                    if not current_q["answer"] and any(k in ans_text for k in correct_keywords[j]):
                                        current_q["answer"] = ['A', 'B', 'C', 'D'][j]

    if current_q and current_q["choices"]:
        questions.append(current_q)
        max_id = max(max_id, current_q["id"])

    text_lines = [p.text for p in doc.paragraphs]
    text_questions = _parse_lines(text_lines)
    for tq in text_questions:
        max_id += 1
        tq["id"] = max_id
        questions.append(tq)

    answer_map = {}
    for table in doc.tables:
        cells = [c.text.strip() for row in table.rows for c in row.cells]
        is_ans = any(re.search(r'ĐÁP\s*ÁN|ĐA\b|MÃ\s*ĐỀ|Mã\s*đề', x, re.IGNORECASE) for x in cells)
        if not is_ans: is_ans = sum(1 for x in cells if x.upper() in 'ABCD') > len(cells) * 0.3
        if is_ans:
            for row in table.rows:
                vals = [c.text.strip() for c in row.cells]
                for i in range(0, len(vals)-1, 2):
                    if re.match(r'^\d+$', vals[i]) and vals[i+1].upper() in 'ABCD':
                        answer_map[int(vals[i])] = vals[i+1].upper()
            rows = list(table.rows)
            if len(rows) >= 2:
                header = [c.text.strip() for c in rows[0].cells]
                for row in rows[1:]:
                    for j, c in enumerate(row.cells):
                        v = c.text.strip().upper()
                        if j < len(header) and re.match(r'^\d+$', header[j]) and v in 'ABCD':
                            answer_map[int(header[j])] = v
            if answer_map: break
            
    for q in questions:
        if not q["answer"] and q["id"] in answer_map:
            q["answer"] = answer_map[q["id"]]

    for q in questions:
        normalize_merged_choices_in_question(q)
    return questions


GROQ_PROMPT = """Bạn là chuyên gia trích xuất đề trắc nghiệm (Toán, Vật lý, Hóa học, Sinh học).

QUY TẮC BẮT BUỘC:
- Mỗi phương án A, B, C, D là một object RIÊNG trong "choices". Không gộp hai phương án vào cùng một "text".
- Trong "text" của từng phương án: không được để lẫn mốc " B. ", " C. ", " D. " — nếu có, tách thành phương án mới.
- "question" chỉ là phần đề; các dòng A. B. C. D. luôn tách vào choices (đủ 4 phương án nếu đề có đủ).
- Giữ ký hiệu toán (căn, phân số, pi, góc) theo đúng nguồn nếu có thể.
- Nếu cuối file có bảng Đáp án / Mã đề (số câu ↔ A/B/C/D), map đúng "answer" theo số câu.
- Nếu có đáp án tô màu/in đậm trong đề, dùng làm "answer"; không thì "" hoặc suy luận nếu chắc chắn.

QUY TẮC STEM (Toán / Lý / Hóa):
- Giữ nguyên subscript: H₂O, CO₂, O₂, N₂, Fe³⁺, SO₄²⁻, CH₃, C₂H₅
- Giữ nguyên superscript: x², x³, n⁺, e⁻, α, β, γ, 10⁶, 10⁻³
- Giữ mũi tên hóa: → (phản ứng), ⇌ (cân bằng), ↑ (khí), ↓ (kết tủa)
- Giữ ký hiệu Hy Lạp: α, β, γ, δ, ω, λ, ν, φ, θ, π, Σ, Δ, Ω, μ, ε, ρ, σ, τ
- Giữ ký hiệu toán: √, ∫, ∑, ∏, ∂, ∇, ∞, ∈, ∉, ⊂, ⊃, ∪, ∩, ≤, ≥, ≈, ≡, ≠, ±, ÷, ×
- Giữ đơn vị: m, s, kg, N, J, W, Pa, Hz, Ω, V, A, T, K, °C, mol, L, km/h, m/s, eV, MeV, GeV
- Giữ công thức hóa: NaCl, H₂SO₄, HCl, NaOH, Ca(OH)₂, NH₄⁺, NO₃⁻, SO₄²⁻, PO₄³⁻, CO₃²⁻
- Giữ vector: →v, →F, →a, vectơ v, vectơ B
- Giữ số mũ vật lý: m², m³, kg·m/s², N·m, J/s, Wb·A
- Phân số: 1/2 → ½, 1/3 → ⅓, 2/3 → ⅔, 3/4 → ¾ (nếu đề gốc dùng fraction thì giữ nguyên)
- Nếu đề là ảnh scanned hoặc text lỗi font: cố gắng suy luận công thức đúng từ ngữ cảnh

QUY TẮC TIẾNG VIỆT:
- Giữ dấu tiếng Việt chuẩn: ầ, ơ, ư, ế, ...
- Không thay đổi nội dung câu hỏi, chỉ tách/sửa format

CHỈ TRẢ VỀ MẢNG JSON THUẦN (không markdown):
[
  {
    "id": 1,
    "question": "...",
    "choices": [{"label": "A", "text": "..."}, {"label": "B", "text": "..."}, {"label": "C", "text": "..."}, {"label": "D", "text": "..."}],
    "answer": "A"
  }
]"""

def parse_with_groq_ai(content: bytes, filetype: str, api_key: str) -> list:
    """Parse PDF/DOCX bằng Groq AI (Llama)."""
    import requests as _req
    try:
        if filetype == "pdf":
            print(f"Đang trích xuất text từ PDF ({len(content)} bytes)...")
            pdf_text = _extract_text_from_pdf(content)
            if not pdf_text.strip():
                print("Không trích xuất được text từ PDF")
                return []
            prompt_text = f"Trích xuất đề trắc nghiệm từ nội dung PDF sau:\n\n{pdf_text[:30000]}"
        else:
            try:
                import docx as _docx
                doc = _docx.Document(io.BytesIO(content))
                text_lines = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        text_lines.append(" | ".join(c.text.replace('\n', ' ').strip() for c in row.cells))
            except Exception:
                text_lines = [
                    q.get("question", "") + " " + " ".join(ch.get("text", "") for ch in q.get("choices", []))
                    for q in parse_docx_without_docx(content)
                ]
            full_text = "\n".join(text_lines)
            if not full_text.strip():
                return []
            prompt_text = f"Trích xuất đề trắc nghiệm từ nội dung Word sau:\n\n{full_text[:30000]}"

        print(f"Đang gửi nội dung lên Groq AI (Llama)...")

        resp = _req.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": "llama-3.3-70b-versatile",
                "messages": [
                    {"role": "system", "content": "Bạn là chuyên gia trích xuất đề trắc nghiệm. Trả lời CHÍNH XÁC mảng JSON thuần, không thêm markdown."},
                    {"role": "user", "content": GROQ_PROMPT + "\n\n" + prompt_text}
                ],
                "temperature": 0.1,
                "max_completion_tokens": 4096
            },
            timeout=120
        )

        if resp.status_code != 200:
            print(f"Groq API error: {resp.status_code} - {resp.text[:200]}")
            return []

        raw = resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")
        raw = re.sub(r'^```json\s*', '', raw)
        raw = re.sub(r'^```\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)

        m = re.search(r'\[.*\]', raw, re.DOTALL)
        if m:
            out = json.loads(m.group())
            for q in out:
                normalize_merged_choices_in_question(q)
            return out
        return []
    except Exception as e:
        print(f"Lỗi khi gọi Groq AI: {e}")
        return []


def _extract_text_from_pdf(content: bytes) -> str:
    """Trích xuất text thuần từ PDF bằng PyPDF2 hoặc pdfplumber."""
    try:
        import io as _io
        try:
            import pypdf
            reader = pypdf.PdfReader(_io.BytesIO(content))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)
        except ImportError:
            pass
        try:
            import pdfplumber
            with pdfplumber.open(_io.BytesIO(content)) as pdf:
                pages = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                return "\n".join(pages)
        except ImportError:
            pass
        return ""
    except Exception:
        return ""



def _has_images_in_pdf(content: bytes) -> bool:
    """Phát hiện PDF có chứa hình ảnh (ảnh scan, công thức, đề scanned) hay không."""
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc:
            images = page.get_images(full=True)
            if images:
                return True
            xrefs = page.get_xobjects()
            for xref in xrefs:
                if xref.get("base_uri") or xref.get("xobj"):
                    return True
        doc.close()
        return False
    except Exception:
        return False


def _extract_images_from_pdf(content: bytes, max_images: int = 50) -> list:
    """
    Trích xuất toàn bộ hình ảnh từ PDF.
    Trả về list of dict: [{"page": int, "index": int, "image": PIL.Image, "width": int, "height": int}]
    """
    extracted = []
    try:
        import fitz
        import io as _io
        doc = fitz.open(stream=content, filetype="pdf")
        for page_num, page in enumerate(doc):
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                if len(extracted) >= max_images:
                    break
                xref = img[0]
                try:
                    base_image = doc.extract_image(xref)
                    img_bytes = base_image.get("image")
                    if not img_bytes:
                        continue
                    ext = base_image.get("ext", "png")
                    width = base_image.get("width", 0)
                    height = base_image.get("height", 0)
                    if width < 100 or height < 50:
                        continue
                    img_pil = None
                    if Image:
                        try:
                            img_pil = Image.open(_io.BytesIO(img_bytes))
                        except Exception:
                            pass
                    extracted.append({
                        "page": page_num + 1,
                        "index": img_index,
                        "image": img_pil,
                        "image_bytes": img_bytes,
                        "width": width,
                        "height": height,
                        "ext": ext,
                    })
                except Exception:
                    continue
            if len(extracted) >= max_images:
                break
        doc.close()
    except Exception as e:
        print(f"[WARN] Lỗi trích xuất ảnh từ PDF: {e}")
    return extracted


def _ocr_image(img_pil, api_key: str) -> str:
    """
    OCR một ảnh bằng Groq Vision API.
    Trả về text trích xuất được từ ảnh.
    """
    try:
        import requests as _req
        import base64 as _b64
        import io as _io

        if not api_key:
            return ""

        if img_pil:
            buf = _io.BytesIO()
            img_pil.save(buf, format="PNG")
            img_bytes = buf.getvalue()
        else:
            return ""

        img_b64 = _b64.b64encode(img_bytes).decode("utf-8")

        payload = {
            "model": "llama-3.2-11b-vision-preview",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/png;base64,{img_b64}"}
                        },
                        {
                            "type": "text",
                            "text": "Đọc toàn bộ nội dung trong ảnh này. Nếu là đề trắc nghiệm, trả về đúng định dạng JSON với các trường: question, choices (mảng A/B/C/D), answer (nếu có). Không thêm markdown."
                        }
                    ]
                }
            ],
            "temperature": 0.1,
            "max_completion_tokens": 2048
        }

        resp = _req.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json=payload,
            timeout=60
        )

        if resp.status_code != 200:
            print(f"[WARN] Groq Vision API error: {resp.status_code}")
            return ""

        raw = resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")
        import re as _re
        m = _re.search(r'\[.*\]', raw, _re.DOTALL)
        if m:
            return m.group()
        return raw.strip()

    except Exception as e:
        print(f"[WARN] OCR image error: {e}")
        return ""


def _ocr_all_images(images: list, api_key: str, batch_size: int = 3) -> list:
    """
    OCR tất cả ảnh trong PDF, trả về list text/json kết quả.
    Xử lý theo batch để tránh rate limit.
    """
    results = []
    for i in range(0, len(images), batch_size):
        batch = images[i:i + batch_size]
        for img_data in batch:
            print(f"  [OCR] Đang xử lý ảnh page {img_data['page']}, index {img_data['index']} "
                  f"({img_data['width']}x{img_data['height']})...")
            ocr_result = _ocr_image(img_data.get("image"), api_key)
            if ocr_result:
                results.append({
                    "page": img_data["page"],
                    "text": ocr_result,
                })
    return results


def _has_text_in_pdf(content: bytes) -> bool:
    """Kiểm tra PDF có text có thể đọc được hay chủ yếu là ảnh (scanned)."""
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        total_text_len = 0
        for page in doc:
            text = page.get_text()
            total_text_len += len(text.strip())
            if total_text_len > 200:
                doc.close()
                return True
        doc.close()
        return total_text_len > 100
    except Exception:
        return False


def parse_pdf_scanned_or_mixed(content: bytes, api_key: str) -> list:
    """
    Luồng OCR cho PDF có ảnh (scanned / ảnh chứa công thức).
    Kết hợp: trích xuất ảnh → OCR Vision → parse bằng AI.
    """
    import fitz
    import re as _re

    print(f"[ScannedPipeline] Phát hiện PDF có ảnh, bắt đầu luồng OCR...")

    pdf_text = _extract_text_from_pdf(content)

    images = _extract_images_from_pdf(content)
    print(f"[ScannedPipeline] Tìm thấy {len(images)} ảnh trong PDF")

    ocr_results = []
    if images and api_key:
        ocr_results = _ocr_all_images(images, api_key)

    combined_text = pdf_text
    if ocr_results:
        ocr_text_parts = [f"[Nội dung từ ảnh trang {r['page']}]: {r['text']}" for r in ocr_results]
        combined_text += "\n\n" + "\n\n".join(ocr_text_parts)

    if not combined_text.strip():
        print("[ScannedPipeline] Không trích xuất được nội dung nào")
        return []

    print(f"[ScannedPipeline] Gửi {len(combined_text)} ký tự lên Groq AI parse...")
    return parse_with_groq_ai_text_only(combined_text, api_key)


def parse_with_groq_ai_text_only(text: str, api_key: str) -> list:
    """Parse text đã trích xuất (từ OCR hoặc PDF thường) bằng Groq AI."""
    try:
        import requests as _req
        import json as _json
        import re as _re

        if not text.strip():
            return []

        prompt_text = f"Trích xuất đề trắc nghiệm từ nội dung sau:\n\n{text[:30000]}"

        resp = _req.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": "llama-3.3-70b-versatile",
                "messages": [
                    {"role": "system", "content": "Bạn là chuyên gia trích xuất đề trắc nghiệm. Trả lời CHÍNH XÁC mảng JSON thuần, không thêm markdown."},
                    {"role": "user", "content": GROQ_PROMPT + "\n\n" + prompt_text}
                ],
                "temperature": 0.1,
                "max_completion_tokens": 4096
            },
            timeout=120
        )

        if resp.status_code != 200:
            print(f"Groq API error: {resp.status_code} - {resp.text[:200]}")
            return []

        raw = resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")
        raw = _re.sub(r'^```json\s*', '', raw)
        raw = _re.sub(r'^```\s*', '', raw)
        raw = _re.sub(r'\s*```$', '', raw)

        m = _re.search(r'\[.*\]', raw, _re.DOTALL)
        if m:
            out = _json.loads(m.group())
            for q in out:
                normalize_merged_choices_in_question(q)
            return out
        return []
    except Exception as e:
        print(f"Lỗi khi gọi Groq AI: {e}")
        return []


def parse_doc_legacy(content: bytes) -> list:
    candidates = ["utf-8", "utf-16", "cp1258", "latin-1"]
    text = ""
    for enc in candidates:
        try:
            text = content.decode(enc)
            break
        except Exception:
            continue
    if not text:
        return []
    lines = [ln.strip("\ufeff\x00 ") for ln in text.replace("\r", "\n").split("\n") if ln.strip()]
    return _parse_lines(lines)


def smart_parse(content: bytes, filename: str, force_ai: bool = False) -> dict:
    fname_lower = filename.lower()
    if fname_lower.endswith(".pdf"):
        filetype = "pdf"
    elif fname_lower.endswith(".doc"):
        filetype = "doc"
    else:
        filetype = "docx"
    cfg = load_config()
    ai_key = cfg.get("groq_key", "").strip()
    ai_ok = bool(ai_key and ai_key != "ĐIỀN_API_KEY_CỦA_BẠN_VÀO_ĐÂY")
    ai_enabled = cfg.get("ai_parse_enabled", True) and ai_ok

    method, error_msg = "normal", ""
    try:
        if filetype == "pdf":
            questions = parse_pdf(content)
        elif filetype == "doc":
            questions = parse_doc_legacy(content)
        else:
            questions = parse_word(content)
    except Exception as e:
        questions = []
        error_msg = str(e)
        print(f"Parse error:\n{traceback.format_exc()}")

    total = len(questions)
    has_ans = sum(1 for q in questions if q.get("answer") in list("ABCD"))
    ans_rate = round(has_ans / total * 100) if total > 0 else 0

    garbled_ratio = _garbled_ratio(questions)
    garbled = garbled_ratio >= 0.2

    if filetype == "pdf" and ai_enabled:
        force_ai = True

    need_ai = force_ai or (ai_enabled and ((total > 0 and ans_rate < 30) or total == 0 or garbled))
    if need_ai and ai_ok:
        if filetype == "pdf":
            has_imgs = _has_images_in_pdf(content)
            has_text = _has_text_in_pdf(content)

            if has_imgs and has_text:
                method = "groq_ai_ocr"
                print(f"[smart_parse] PDF hỗn hợp (text + images), dùng OCR pipeline")
                try:
                    ai_questions = parse_pdf_scanned_or_mixed(content, ai_key)
                except Exception as e:
                    print(f"OCR pipeline error: {e}")
                    ai_questions = []
            elif has_imgs:
                method = "groq_ai_ocr"
                print(f"[smart_parse] PDF scanned/ảnh, dùng OCR pipeline")
                try:
                    ai_questions = parse_pdf_scanned_or_mixed(content, ai_key)
                except Exception as e:
                    print(f"OCR pipeline error: {e}")
                    ai_questions = []
            else:
                method = "groq_ai"
                print(f"[smart_parse] PDF text thuần, dùng AI parser")
                try:
                    ai_questions = parse_with_groq_ai(content, filetype, ai_key)
                except Exception as e:
                    print(f"Groq AI error: {e}")
                    ai_questions = []
        else:
            method = "groq_ai"
            try:
                ai_questions = parse_with_groq_ai(content, filetype, ai_key)
            except Exception as e:
                print(f"Groq AI error: {e}")
                ai_questions = []

        if ai_questions:
            for i, q in enumerate(ai_questions):
                q["id"] = i + 1
            questions = ai_questions

        total = len(questions)
        has_ans = sum(1 for q in questions if q.get("answer") in list("ABCD"))
        ans_rate = round(has_ans / total * 100) if total > 0 else 0

    warning = ""
    if garbled and method != "groq_ai" and method != "groq_ai_ocr":
        warning = "PDF có dấu hiệu lỗi font/công thức; nên bật AI parse để tăng độ chính xác."

    return {
        "questions": questions,
        "total": total,
        "ans_rate": ans_rate,
        "method": method,
        "ai_available": ai_ok,
        "error": error_msg,
        "warning": warning,
        "garbled_ratio": round(garbled_ratio, 4),
        "has_images": _has_images_in_pdf(content) if filetype == "pdf" else False,
    }



class RegisterBody(BaseModel):
    email: str
    password: str
    role: str = "user"


class LoginBody(BaseModel):
    email: str
    password: str


class RefreshBody(BaseModel):
    refresh_token: str


class ProfileBody(BaseModel):
    display_name: str = ""
    avatar_url: str = ""


def _auth_payload(user: dict) -> dict:
    access_token = create_access_token(user)
    refresh_token = create_refresh_token(user)
    return {
        "user": user,
        "access_token": access_token,
        "refresh_token": refresh_token,
        "token_type": "bearer",
    }


def require_admin(user: dict) -> None:
    if not user or (user.get("role") or "user") != "admin":
        raise HTTPException(403, "Chỉ admin mới có quyền thực hiện.")


@app.post("/auth/register")
def register(body: RegisterBody):
    email = (body.email or "").strip().lower()
    if not email:
        raise HTTPException(400, "Email không hợp lệ")
    uid = f"u_{uuid.uuid4().hex[:20]}"
    try:
        role = (body.role or "user").strip().lower()
        user = repo.register_user(uid, email, body.password, role=role)
        return {"message": "Đăng ký thành công", **_auth_payload(user)}
    except ValueError as e:
        raise HTTPException(400, str(e))


@app.post("/auth/login")
def login(body: LoginBody):
    user = repo.authenticate_user(body.email, body.password)
    if not user:
        raise HTTPException(401, "Sai email hoặc mật khẩu")
    return {"message": "Đăng nhập thành công", **_auth_payload(user)}


@app.post("/auth/refresh")
def refresh_token(body: RefreshBody):
    token = (body.refresh_token or "").strip()
    if not token:
        raise HTTPException(400, "Thiếu refresh_token")
    payload = decode_refresh_token(token)
    uid = (payload.get("sub") or "").strip()
    if not uid:
        raise HTTPException(401, "Token không hợp lệ")
    user = repo.get_user_by_uid(uid)
    if not user:
        raise HTTPException(401, "Người dùng không tồn tại")

    exp = payload.get("exp")
    exp_dt = datetime.utcfromtimestamp(int(exp)) if exp else None
    repo.revoke_token(token, "refresh", uid, exp_dt)

    return {"message": "Làm mới token thành công", **_auth_payload(user)}


@app.post("/auth/logout")
def logout(body: RefreshBody, authorization: str = Header(default="")):
    auth = (authorization or "").strip()
    if not auth.lower().startswith("bearer "):
        raise HTTPException(401, "Thiếu Bearer token")
    access_token = auth[7:].strip()
    if not access_token:
        raise HTTPException(401, "Thiếu Bearer token")

    access_payload = decode_access_token(access_token)
    access_exp = access_payload.get("exp")
    access_exp_dt = datetime.utcfromtimestamp(int(access_exp)) if access_exp else None
    uid = (access_payload.get("sub") or "").strip() or None
    repo.revoke_token(access_token, "access", uid, access_exp_dt)

    refresh_token_str = (body.refresh_token or "").strip()
    if refresh_token_str:
        try:
            refresh_payload = decode_refresh_token(refresh_token_str)
            refresh_exp = refresh_payload.get("exp")
            refresh_exp_dt = datetime.utcfromtimestamp(int(refresh_exp)) if refresh_exp else None
            refresh_uid = (refresh_payload.get("sub") or "").strip() or uid
            repo.revoke_token(refresh_token_str, "refresh", refresh_uid, refresh_exp_dt)
        except HTTPException:
            pass

    return {"message": "Đăng xuất thành công"}


@app.get("/auth/me")
def me(authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    return {"user": user}


@app.put("/auth/profile")
def update_profile(body: ProfileBody, authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    display_name = (body.display_name or "").strip()
    avatar_url = (body.avatar_url or "").strip()
    updated = repo.update_user_profile(user["uid"], display_name=display_name, avatar_url=avatar_url)
    if not updated:
        raise HTTPException(404, "Không tìm thấy người dùng")
    return {"message": "Đã lưu hồ sơ", "user": updated}


@app.get("/admin/users")
def admin_list_users(authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    require_admin(user)
    return {"users": repo.list_users()}


class UpdateRoleBody(BaseModel):
    role: str


@app.put("/admin/users/{uid}/role")
def admin_update_role(uid: str, body: UpdateRoleBody, authorization: str = Header(default="")):
    admin_user = get_auth_user(authorization)
    require_admin(admin_user)
    role = (body.role or "").strip().lower()
    if role not in ("user", "admin"):
        raise HTTPException(400, "Role chỉ nhận user/admin")
    ok = repo.set_user_role(uid, role)
    if not ok:
        raise HTTPException(404, "Không tìm thấy user")
    return {"message": "Đã cập nhật role", "uid": uid, "role": role}


@app.get("/config")
def get_config():
    ok = has_valid_key()
    cfg = load_config()
    return {
        "has_api_key": ok,
        "has_key": ok,
        "ai_enabled": cfg.get("ai_parse_enabled", True)
    }

class ConfigBody(BaseModel):
    gemini_key: str = ""
    ai_parse_enabled: bool = True

@app.post("/config")
def update_config(body: ConfigBody):
    repo.set_ai_parse_enabled(body.ai_parse_enabled)
    return {"message": "Đã lưu cài đặt", "has_key": has_valid_key()}

class AIConfigBody(BaseModel):
    enabled: bool = True

class AIFillConfigBody(BaseModel):
    enabled: bool = False

@app.post("/config/ai")
def update_ai_config(body: AIConfigBody):
    repo.set_ai_parse_enabled(body.enabled)
    return {"message": "Đã lưu cài đặt AI", "ai_enabled": body.enabled, "has_key": has_valid_key()}

@app.post("/config/ai-fill")
def update_ai_fill_config(body: AIFillConfigBody):
    """Toggle AI Fill feature on/off (independent of Gemini key for upload)."""
    repo.set_config_value("ai_fill_enabled", body.enabled)
    return {"message": "Đã lưu", "ai_fill_enabled": body.enabled}

@app.delete("/config/key")
def delete_api_key():
    repo.set_ai_parse_enabled(True)
    return {"message": "Đã reset config"}

@app.get("/stats")
def get_stats(authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    index = repo.get_files_index(uid)
    agg = repo.get_stats_aggregate(uid)
    ok = has_valid_key()
    cfg = load_config()
    return {
        "total_questions": agg["total_questions"],
        "with_answer": agg["with_answer"],
        "total_sessions": agg["total_sessions"],
        "avg_score": agg["avg_score"],
        "best_score": agg["best_score"],
        "ai_available": ok,
        "ai_enabled": cfg.get("ai_parse_enabled", True),
        "ai_fill_enabled": cfg.get("ai_fill_enabled", False),
        "files": [{"id": fid, "name": f["name"], "count": f["count"], "with_answer": f["with_answer"], "uploaded_at": f["uploaded_at"], "parse_method": f.get("parse_method", "normal")} for fid, f in index.items()]
    }

@app.post("/upload")
async def upload_file(file: UploadFile = File(...), authorization: str = Header(default=""), x_force_ai: str = Header(default="false")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    content = await file.read()
    fname = file.filename or "file"

    if not (fname.lower().endswith(".docx") or fname.lower().endswith(".doc") or fname.lower().endswith(".pdf")):
        raise HTTPException(400, "Chỉ hỗ trợ .doc, .docx và .pdf")

    force_ai = x_force_ai.lower() == "true"
    try:
        result = smart_parse(content, fname, force_ai=force_ai)
    except HTTPException: raise
    except Exception as e: raise HTTPException(500, f"Lỗi xử lý: {e}")

    questions = result["questions"]
    if not questions:
        err = result.get("error", "")
        raise HTTPException(400, f"Không tìm thấy câu hỏi.{' Chi tiết: '+err if err else ' Thử bật AI Vision.'}")

    if fname.lower().endswith('.pdf') and result.get('ai_available') and result.get('method') != 'groq_ai':
        raise HTTPException(400, "PDF bắt buộc parse bằng AI khi AI khả dụng. Vui lòng bật Force AI và upload lại.")

    if fname.lower().endswith('.pdf'):
        g_ratio = float(result.get('garbled_ratio') or 0)
        ans_rate = float(result.get('ans_rate') or 0)
        method = result.get('method') or 'unknown'

        if method != 'groq_ai' and (ans_rate < 35 or g_ratio >= 0.12):
            raise HTTPException(
                400,
                f"PDF parse thường chất lượng thấp (đáp án={ans_rate:.0f}%, garbled={g_ratio:.2f}). Hãy bật AI Vision rồi upload lại."
            )

        if method == 'groq_ai' and g_ratio >= 0.30:
            raise HTTPException(
                400,
                f"PDF parse AI vẫn bị lỗi ký tự nặng (garbled={g_ratio:.2f}). Hard-stop để tránh lưu đề lỗi."
            )


    base_name = re.sub(r'\.(docx|doc|pdf)$', '', fname, flags=re.IGNORECASE)
    file_id = re.sub(r'[^\w\-]', '_', base_name)[:50]

    repo.ensure_user(uid)
    normalized_questions = []
    for i, q in enumerate(questions, start=1):
        q["id"] = i
        q = enrich_question_rich_fields(q)
        normalized_questions.append(q)

    uploaded_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    repo.replace_file_questions(
        uid, file_id, normalized_questions, base_name, fname, uploaded_at, result["method"]
    )
    has_ans = sum(1 for q in normalized_questions if q.get("answer") in list("ABCD"))

    return {"file_id": file_id, "name": base_name, "parsed": result["total"], "added": len(normalized_questions), "total_in_file": len(normalized_questions), "with_answer": has_ans, "ans_rate": result["ans_rate"], "parse_method": result["method"], "ai_available": result["ai_available"], "message": "Upload thành công (đã ghi đè đề cũ)"}

@app.delete("/files/{file_id}")
def delete_file(file_id: str, authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    if not repo.file_exists(uid, file_id):
        raise HTTPException(404, "Không tìm thấy file")
    repo.delete_questions_file(uid, file_id)
    return {"message": "Đã xóa"}

@app.get("/quiz/start")
def start_quiz(num: int = 10, file_id: str = "", authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    if file_id:
        all_qs = repo.get_questions_json(uid, file_id)
    else:
        all_qs = repo.get_all_questions(uid)

    valid = [q for q in all_qs if len(q.get("choices", [])) >= 2]
    if not valid: raise HTTPException(404, "Không có câu hỏi")

    selected = random.sample(valid, min(num, len(valid)))
    sid = f"s{random.randint(100000, 999999)}"
    quiz = []
    choice_labels = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    for i, q in enumerate(selected):
        orig = q["choices"]
        n = len(orig)
        order = list(range(n))
        random.shuffle(order)
        choices = []
        for pos, orig_i in enumerate(order):
            lab = choice_labels[pos] if pos < len(choice_labels) else str(pos + 1)
            choices.append({"label": lab, "text": orig[orig_i]["text"]})
        ans_lbl = (q.get("answer") or "").strip().upper()
        new_correct = ""
        if ans_lbl:
            try:
                orig_i = next(
                    j for j, c in enumerate(orig)
                    if (c.get("label") or "").strip().upper() == ans_lbl
                )
            except StopIteration:
                orig_i = None
            if orig_i is not None:
                pos = order.index(orig_i)
                new_correct = choice_labels[pos] if pos < len(choice_labels) else str(pos + 1)
        quiz.append({"id": i, "question": q["question"], "choices": choices, "_correct": new_correct})
    repo.save_quiz_session(uid, sid, file_id or "", quiz)
    return {"session_id": sid, "total": len(quiz), "file_id": file_id, "questions": [{"id": q["id"], "question": q["question"], "choices": q["choices"]} for q in quiz]}

class SubmitBody(BaseModel):
    session_id: str
    answers: dict
    time_taken: int = 0
    file_id: str = ""
    anti_cheat: dict = {}

@app.post("/quiz/submit")
def submit_quiz(body: SubmitBody, authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    quiz = repo.get_quiz_session(body.session_id)
    if not quiz:
        raise HTTPException(404, "Session không tồn tại")
    details = []
    for q in quiz:
        user = body.answers.get(str(q["id"]), "")
        ok = bool(user and user == q["_correct"])
        details.append({"id": q["id"], "question": q["question"], "user": user, "correct": q["_correct"], "ok": ok, "choices": q["choices"]})
    score = sum(1 for d in details if d["ok"])
    pct = round(score / len(quiz) * 100)
    wrong_q = [d["question"][:60] for d in details if not d["ok"]][:5]
    ac_in = body.anti_cheat or {}
    anti_cheat = {
        "blur_count": int(ac_in.get("blur_count") or ac_in.get("tab_switches") or 0),
        "penalty_base_sec": int(ac_in.get("penalty_base_sec") or 15),
    }
    review_details = []
    for d in details:
        review_details.append(
            {
                "id": d["id"],
                "question": d["question"],
                "user": d["user"],
                "correct": d["correct"],
                "ok": d["ok"],
                "choices": d.get("choices") or [],
            }
        )

    repo.append_history(
        uid,
        score,
        len(quiz),
        pct,
        body.time_taken,
        body.file_id or "all",
        wrong_q,
        anti_cheat=anti_cheat,
        review_details=review_details,
    )
    repo.delete_quiz_session(body.session_id)
    return {"score": score, "total": len(quiz), "percent": pct, "details": details}

@app.get("/history")
def get_history(authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    return repo.get_history_list(user["uid"])

@app.delete("/history/clear")
def clear_history(authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    repo.clear_history(user["uid"])
    return {"message": "Đã xóa lịch sử"}

@app.get("/files/{file_id}/questions")
def get_file_questions(file_id: str, authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    if not repo.file_exists(uid, file_id):
        raise HTTPException(404, "Không tìm thấy file")
    questions = repo.get_questions_json(uid, file_id)
    return {"questions": questions, "total": len(questions)}

class QuestionUpdateBody(BaseModel):
    question: str
    choices: list
    answer: str = ""


class QuestionReviewBody(BaseModel):
    reviewed: bool = True

@app.put("/files/{file_id}/questions/{q_id}")
def update_question(file_id: str, q_id: int, body: QuestionUpdateBody, authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    if not repo.file_exists(uid, file_id):
        raise HTTPException(404, "Không tìm thấy file")
    choices, answer = _validate_question_choices(body.choices, body.answer)
    updated = repo.update_question_row(uid, file_id, q_id, body.question, choices, answer)
    if not updated:
        raise HTTPException(404, "Không tìm thấy câu hỏi")
    return {"message": "Đã cập nhật", "question": updated}

@app.put("/files/{file_id}/questions/{q_id}/review")
def set_question_review(file_id: str, q_id: int, body: QuestionReviewBody, authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    if not repo.file_exists(uid, file_id):
        raise HTTPException(404, "Không tìm thấy file")
    updated = repo.set_question_reviewed(uid, file_id, q_id, bool(body.reviewed))
    if not updated:
        raise HTTPException(404, "Không tìm thấy câu hỏi")
    return {"message": "Đã cập nhật review", "question": updated}


@app.delete("/files/{file_id}/questions/{q_id}")
def delete_question(file_id: str, q_id: int, authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    if not repo.file_exists(uid, file_id):
        raise HTTPException(404, "Không tìm thấy file")
    if not repo.delete_question_by_qid(uid, file_id, q_id):
        raise HTTPException(404, "Không tìm thấy câu hỏi")
    return {"message": "Đã xóa câu hỏi"}

class NewQuestionBody(BaseModel):
    question: str
    choices: list
    answer: str = ""

@app.post("/files/{file_id}/questions")
def add_question(file_id: str, body: NewQuestionBody, authorization: str = Header(default="")):
    user = get_auth_user(authorization)
    uid = user["uid"]
    if not repo.file_exists(uid, file_id):
        raise HTTPException(404, "Không tìm thấy file")
    choices, answer = _validate_question_choices(body.choices, body.answer)
    questions = repo.get_questions_json(uid, file_id)
    max_id = max((q.get("id", 0) for q in questions), default=0)
    new_q = {
        "id": max_id + 1,
        "question": body.question,
        "choices": choices,
        "answer": answer,
        "explanation": "",
    }
    repo.insert_question(uid, file_id, new_q)
    return {"message": "Đã thêm câu hỏi", "question": new_q}



class AIAnswerQuestionBody(BaseModel):
    question: str
    choices: list
    subject: str = ""
    context: str = ""

class AIAnswerFileBody(BaseModel):
    file_id: str
    use_rag: bool = True
    force_llm: bool = False
    subject: str = ""


@app.post("/ai/answer-question")
async def ai_answer_question(
    body: AIAnswerQuestionBody,
    authorization: str = Header(default="")
):
    """
    Trả lời 1 câu hỏi trắc nghiệm bằng AI (MRC).
    Input: question + choices [+ context từ RAG]
    Output: answer + confidence + explanation
    """
    _init_ai_services()

    if not _is_ai_pipeline_ready() or not _ai_pipeline:
        raise HTTPException(503, "AI Pipeline chưa sẵn sàng. Vui lòng kiểm tra GROQ_API_KEY.")

    user = get_auth_user(authorization)

    result = _ai_pipeline.process_question(
        question_data={
            "question": body.question,
            "choices": body.choices,
            "subject": body.subject
        },
        subject=body.subject,
        use_rag=False,
        force_llm=True
    )

    _log_ai_call(
        user_id=user["uid"],
        question_id=None,
        file_id="",
        model=_mrc_service.model_name if _mrc_service else "unknown",
        prompt="",
        response=json.dumps(result, ensure_ascii=False),
        success=not result.get("error", False)
    )

    return result


@app.post("/ai/answer-file")
async def ai_answer_file(
    body: AIAnswerFileBody,
    authorization: str = Header(default="")
):
    """
    Tự động trả lời TẤT CẢ câu hỏi trong 1 file.
    - RAG: tìm câu tương tự trong vector DB
    - LLM: gọi Gemini nếu không tìm thấy
    - Cache: lưu kết quả vào DB
    """
    _init_ai_services()

    if not _is_ai_pipeline_ready() or not _ai_pipeline:
        raise HTTPException(503, "AI Pipeline chưa sẵn sàng. Vui lòng kiểm tra GROQ_API_KEY.")

    user = get_auth_user(authorization)
    uid = user["uid"]

    if not repo.file_exists(uid, body.file_id):
        raise HTTPException(404, "Không tìm thấy file")

    questions = repo.get_questions_json(uid, body.file_id)
    if not questions:
        raise HTTPException(404, "File không có câu hỏi nào")

    subject = body.subject

    results = []
    stats = {
        "from_vector_db": 0,
        "from_llm": 0,
        "existing": 0,
        "errors": 0,
        "total_confidence": 0.0
    }

    questions_to_fill = [
        (q, bool(q.get("answer") and str(q.get("answer")).strip() in ["A","B","C","D"]))
        for q in questions
    ]
    need_fill = [(q, had_ans) for q, had_ans in questions_to_fill if not had_ans]
    total_to_fill = len(need_fill)

    for idx, (q, has_existing_answer) in enumerate(questions_to_fill, 1):
        if has_existing_answer:
            results.append({
                "question_id": q.get("id"),
                "question": q.get("question", ""),
                "answer": q.get("answer", ""),
                "confidence": 1.0,
                "explanation": "Đáp án đã có sẵn",
                "source": "existing",
                "similar_questions": [],
                "saved": True,
            })
            stats["existing"] += 1
            continue

        result = _ai_pipeline.process_question(
            question_data={
                "id": q.get("id"),
                "question": q.get("question", ""),
                "choices": q.get("choices", []),
                "file_id": body.file_id,
                "subject": subject
            },
            subject=subject,
            use_rag=body.use_rag,
            force_llm=body.force_llm
        )
        result["question_id"] = q.get("id")
        result["question"] = q.get("question", "")
        results.append(result)

        src = result.get("source", "")
        if src in ("llm", "vector_db") and result.get("answer"):
            try:
                saved = repo.update_question_row(
                    uid=uid,
                    file_id=body.file_id,
                    q_id=int(q.get("id", 0)),
                    question=q.get("question", ""),
                    choices=q.get("choices", []),
                    answer=result["answer"]
                )
                result["saved"] = saved is not None
                if saved is None:
                    logger.warning(f"update_question_row returned None for qid={q.get('id')}")
            except Exception as save_err:
                logger.error(f"Khong luu duoc answer cho qid={q.get('id')}: {save_err}")
                result["saved"] = False
        else:
            result["saved"] = False

        if src == "vector_db":
            stats["from_vector_db"] += 1
        elif src == "llm":
            stats["from_llm"] += 1
        elif src == "existing":
            stats["existing"] += 1
        else:
            stats["errors"] += 1

        stats["total_confidence"] += result.get("confidence", 0)

        is_last = (idx == len(questions_to_fill) - 1)
        if not is_last:
            import time
            time.sleep(0.15)

    if results:
        stats["avg_confidence"] = round(stats["total_confidence"] / len(results), 4)

    stats["filled"] = stats["from_vector_db"] + stats["from_llm"]

    return {
        "file_id": body.file_id,
        "total": len(results),
        "subject": subject,
        "results": results,
        "summary": stats,
        "rag_enabled": body.use_rag,
        "vector_db_total": _vector_store.count() if _vector_store else 0
    }


@app.get("/ai/similar-questions/{question_id}")
async def get_similar_questions(
    question_id: int,
    file_id: str,
    top_k: int = 5,
    subject: str = "",
    authorization: str = Header(default="")
):
    """
    Tìm câu hỏi tương tự trong vector DB.
    Dùng để tham khảo trước khi làm quiz.
    """
    _init_ai_services()

    if not _is_ai_pipeline_ready() or not _rag_service:
        raise HTTPException(503, "AI Pipeline chưa sẵn sàng")

    user = get_auth_user(authorization)
    uid = user["uid"]

    if not repo.file_exists(uid, file_id):
        raise HTTPException(404, "Không tìm thấy file")

    questions = repo.get_questions_json(uid, file_id)
    target_q = next((q for q in questions if q["id"] == question_id), None)
    if not target_q:
        raise HTTPException(404, "Không tìm thấy câu hỏi")

    similar = _rag_service.retrieve_similar(
        question=target_q["question"],
        choices=target_q["choices"],
        top_k=top_k,
        subject=subject if subject else None
    )

    return {
        "question": target_q,
        "similar_questions": similar,
        "total_found": len(similar)
    }


@app.post("/ai/answer-quiz")
async def ai_answer_quiz(
    session_id: str,
    authorization: str = Header(default="")
):
    """
    Tự động trả lời toàn bộ quiz session (chỉ lấy kết quả từ AI).
    """
    _init_ai_services()

    if not _is_ai_pipeline_ready() or not _ai_pipeline:
        raise HTTPException(503, "AI Pipeline chưa sẵn sàng")

    user = get_auth_user(authorization)
    uid = user["uid"]

    quiz = repo.get_quiz_session(session_id)
    if not quiz:
        raise HTTPException(404, "Session không tồn tại")

    ai_answers = {}
    explanations = {}

    for q in quiz:
        qid = str(q["id"])
        result = _ai_pipeline.process_question(
            question_data={
                "question": q["question"],
                "choices": q["choices"]
            },
            use_rag=True,
            force_llm=False
        )
        ai_answers[qid] = result.get("answer", "")
        explanations[qid] = {
            "explanation": result.get("explanation", ""),
            "confidence": result.get("confidence", 0),
            "source": result.get("source", "")
        }

    return {
        "session_id": session_id,
        "ai_answers": ai_answers,
        "explanations": explanations,
        "total": len(quiz)
    }


@app.get("/ai/stats")
async def get_ai_stats(authorization: str = Header(default="")):
    """
    Thống kê AI: cache hit rate, vector DB size, etc.
    """
    _init_ai_services()

    user = get_auth_user(authorization)
    uid = user["uid"]

    stats = {}
    if _vector_store:
        vstats = _vector_store.get_stats()
        stats.update(vstats)

    try:
        from db import get_connection
        conn = get_connection()
        cur = conn.cursor(dictionary=True)
        cur.execute(
            "SELECT COUNT(*) as cnt FROM ai_cache WHERE user_id = %s",
            (uid,)
        )
        row = cur.fetchone()
        stats["user_cache_count"] = row["cnt"] if row else 0

        cur.execute(
            """
            SELECT COUNT(*) as cnt, SUM(total_tokens) as tokens, SUM(cost_estimate) as cost
            FROM ai_llm_logs
            WHERE user_id = %s AND created_at >= DATE_SUB(NOW(), INTERVAL 1 DAY)
            """,
            (uid,)
        )
        row = cur.fetchone()
        stats["llm_calls_24h"] = row["cnt"] if row else 0
        stats["tokens_24h"] = row["tokens"] if row and row["tokens"] else 0
        stats["cost_24h_usd"] = round(row["cost"] if row and row["cost"] else 0, 6)
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Error getting AI stats: {e}")

    return {
        "user_id": uid,
        "vector_db": stats,
        "ai_available": _is_ai_pipeline_ready(),
        "mrc_model": _mrc_service.model_name if _mrc_service else None
    }


@app.post("/ai/cache/clear")
async def clear_ai_cache(
    file_id: str = "",
    subject: str = "",
    authorization: str = Header(default="")
):
    """
    Xóa cache AI (cho 1 file hoặc 1 môn).
    """
    _init_ai_services()

    user = get_auth_user(authorization)
    uid = user["uid"]

    deleted = 0
    try:
        from db import get_connection
        conn = get_connection()
        cur = conn.cursor()

        if file_id:
            if _vector_store:
                _vector_store.delete_by_file(file_id)
            cur.execute("DELETE FROM ai_cache WHERE user_id = %s AND file_id = %s", (uid, file_id))
            deleted = cur.rowcount
        elif subject:
            if _vector_store:
                _vector_store.delete_by_subject(subject)
            cur.execute("DELETE FROM ai_cache WHERE user_id = %s AND subject = %s", (uid, subject))
            deleted = cur.rowcount
        else:
            cur.execute("DELETE FROM ai_cache WHERE user_id = %s", (uid,))
            deleted = cur.rowcount

        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Error clearing AI cache: {e}")
        raise HTTPException(500, str(e))

    return {"message": f"Đã xóa {deleted} cache entries"}



def _log_ai_call(
    user_id: str,
    question_id: int = None,
    file_id: str = "",
    model: str = "",
    prompt: str = "",
    response: str = "",
    prompt_tokens: int = 0,
    response_tokens: int = 0,
    success: bool = True,
    error_message: str = ""
):
    """Log AI LLM call vào DB."""
    try:
        from db import get_connection
        conn = get_connection()
        cur = conn.cursor()
        cur.execute(
            """
            INSERT INTO ai_llm_logs
            (user_id, question_id, file_id, model, prompt_text, response_text,
             prompt_tokens, response_tokens, total_tokens, cost_estimate, success, error_message)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """,
            (
                user_id, question_id, file_id, model,
                prompt[:5000], response[:10000],
                prompt_tokens, response_tokens,
                prompt_tokens + response_tokens,
                0.0,
                1 if success else 0,
                error_message
            )
        )
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Failed to log AI call: {e}")



class ScienceMetadataBody(BaseModel):
    question: str = ""
    choices: list = []
    subject_override: str = ""


@app.post("/science/metadata")
async def science_metadata(
    body: ScienceMetadataBody,
    authorization: str = Header(default="")
):
    """
    Phân tích metadata khoa học cho câu hỏi:
    - Phát hiện môn học (math/physics/chemistry/biology)
    - Phân loại công thức toán (fraction, sqrt, exponent, integral, limit, ...)
    - Phát hiện công thức hóa học, phản ứng
    - Phát hiện đơn vị vật lý, vector, ký hiệu khoa học
    - Chuẩn hóa subscript/superscript
    """
    if not _SCIENCE_PARSER_OK:
        return {"error": "Science parser not available", "ok": False}

    q = {
        "question": body.question,
        "choices": body.choices,
    }

    if body.subject_override:
        q["subject"] = body.subject_override

    enriched = enrich_science_fields(q)

    return {
        "ok": True,
        "subject": enriched.get("subject", "unknown"),
        "subject_confidence": enriched.get("subject_confidence", 0.0),
        "subject_details": enriched.get("subject_details", {}),
        "is_stem": enriched.get("is_stem", False),
        "math_types": enriched.get("math_types", []),
        "math_formula_count": enriched.get("math_formula_count", 0),
        "physics": {
            "has_vector": enriched.get("physics", {}).get("has_vector", False),
            "has_units": enriched.get("physics", {}).get("has_units", False),
            "symbols": enriched.get("physics", {}).get("symbols", []),
            "formulas": enriched.get("physics", {}).get("formulas", []),
        },
        "chemistry": {
            "has_reaction": enriched.get("chemistry", {}).get("has_reaction", False),
            "has_formula": enriched.get("chemistry", {}).get("has_formula", False),
            "formulas": enriched.get("chemistry", {}).get("formulas", [])[:10],
        },
        "formula_count": enriched.get("formula_count", 0),
        "normalized_question": enriched.get("_normalized_question", body.question),
        "normalized_choices": enriched.get("_normalized_choices", body.choices),
    }


@app.post("/science/normalize")
async def science_normalize(
    body: ScienceMetadataBody,
    authorization: str = Header(default="")
):
    """
    Chuẩn hóa công thức khoa học trong text:
    - H2O → H₂O (subscript hóa học)
    - 1/2 → ½ (phân số)
    - --> → → (mũi tên)
    - m/s → chuẩn hóa đơn vị
    - x^2 → x² (superscript)
    """
    if not _SCIENCE_PARSER_OK:
        raise HTTPException(503, "Science parser not available")

    text = body.question or ""
    text = normalize_chemical_formula(normalize_physics_units(text))

    normalized_choices = []
    for c in body.choices or []:
        txt = str((c or {}).get("text") or "").strip()
        lb = (c or {}).get("label", "").strip().upper()
        txt = normalize_chemical_formula(normalize_physics_units(txt))
        normalized_choices.append({"label": lb, "text": txt})

    math_types = classify_math_expression(text)
    physics = detect_physics_elements(text)
    chemistry = detect_chemistry_elements(text)
    subject, subj_conf, _ = detect_subject(f"{text} {' '.join(normalized_choices)}")

    return {
        "normalized_text": text,
        "normalized_choices": normalized_choices,
        "math_types": math_types,
        "physics_units_found": physics.get("has_units", False),
        "chemistry_formulas_found": chemistry.get("has_formula", False),
        "detected_subject": subject,
        "subject_confidence": subj_conf,
    }
