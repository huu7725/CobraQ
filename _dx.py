import docx, io, json

with open(r"C:\Users\TUF\Downloads\thuvienhoclieu.com-De-cuong-on-tap-cuoi-HK2-Lich-su-12-Canh-dieu-24-25.docx", "rb") as f:
    content = f.read()

doc = docx.Document(io.BytesIO(content))

# Save paragraphs and tables to JSON
result = {
    "paragraph_count": len(doc.paragraphs),
    "table_count": len(doc.tables),
    "paragraphs": [p.text for p in doc.paragraphs[:50] if p.text.strip()],
    "tables": []
}

for ti, t in enumerate(doc.tables[:5]):
    rows_data = []
    for row in t.rows[:6]:
        cells = [c.text for c in row.cells]
        rows_data.append(cells)
    result["tables"].append({"rows": rows_data})

with open("d:/CobraQ/_dx.json", "w", encoding="utf-8") as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print("Saved to _dx.json")
